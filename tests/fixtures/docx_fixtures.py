from __future__ import annotations

import base64
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FIXTURE_DIR = Path(__file__).resolve().parent / "docx"

TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def ensure_docx_fixtures() -> dict[str, Path]:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    builders = {
        "style_inheritance_ok": _build_style_inheritance_ok,
        "cover_image_without_caption": _build_cover_image_without_caption,
        "body_image_without_caption": _build_body_image_without_caption,
        "hand_toc_bad_caption_numbers": _build_hand_toc_bad_caption_numbers,
        "commented_highlighted": _build_commented_highlighted,
        "tracked_changes": _build_tracked_changes,
        "page_numbering_sections": _build_page_numbering_sections,
        "advanced_manual_review": _build_advanced_manual_review,
        "template_margin_variants": _build_template_margin_variants,
    }

    paths: dict[str, Path] = {}
    for name, builder in builders.items():
        path = FIXTURE_DIR / f"{name}.docx"
        builder(path)
        paths[name] = path
    return paths


def _build_style_inheritance_ok(path: Path) -> None:
    doc = Document()
    _set_a4_margins(doc)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_before = Pt(6)
    normal.paragraph_format.space_after = Pt(6)

    paragraph = doc.add_paragraph(style="Normal")
    paragraph.add_run(
        "Đây là đoạn văn nội dung đủ dài để kiểm tra kế thừa style trong Word. "
        "Run không đặt trực tiếp cỡ chữ nhưng style Normal đã đặt Times New Roman 13pt."
    )
    paragraph.add_run(" Từ khóa kế thừa style.")
    doc.save(path)


def _build_cover_image_without_caption(path: Path) -> None:
    doc = Document()
    _set_a4_margins(doc)
    title = doc.add_paragraph("TRƯỜNG ĐẠI HỌC MẪU")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_picture_to_paragraph(doc, logo_paragraph, width_cm=2.0)
    doc.add_paragraph("ĐỒ ÁN TỐT NGHIỆP")
    doc.save(path)


def _build_body_image_without_caption(path: Path) -> None:
    doc = Document()
    _set_a4_margins(doc)
    for index in range(1, 27):
        doc.add_paragraph(
            f"Đoạn mở đầu số {index} dùng để đưa ảnh ra khỏi vùng bìa, tránh nhầm logo hoặc ảnh sinh viên."
        )
    image_paragraph = doc.add_paragraph()
    _add_picture_to_paragraph(doc, image_paragraph, width_cm=4.0)
    doc.add_paragraph(
        "Đoạn sau ảnh không phải caption, vì vậy ảnh trong phần nội dung cần được báo thiếu caption gần vị trí ảnh."
    )
    doc.save(path)


def _build_hand_toc_bad_caption_numbers(path: Path) -> None:
    doc = Document()
    _set_a4_margins(doc)
    doc.add_paragraph("MỤC LỤC")
    doc.add_paragraph("Chương 1. Tổng quan 1")
    doc.add_paragraph("DANH MỤC BẢNG")
    doc.add_paragraph("Bảng 3.28. Thống kê lỗi ........ 10")
    doc.add_paragraph("Bảng 3.28. Thống kê lỗi ........ 11")
    doc.add_paragraph("Chương 1.")
    doc.add_paragraph("TỔNG QUAN")
    doc.add_paragraph("Hình 2.2.3.3. Số hình bị rối")
    doc.add_paragraph(
        "Đoạn nội dung đủ dài để analyzer coi là thân bài và giữ ngữ cảnh chương gần nhất."
    )
    doc.save(path)


def _build_commented_highlighted(path: Path) -> None:
    doc = Document()
    _set_a4_margins(doc)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(
        "Đoạn văn có comment và highlight để kiểm tra file fixed phải sạch nhưng giữ nguyên nội dung."
    )
    run.font.name = "Arial"
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if hasattr(doc, "add_comment"):
        doc.add_comment([run], text="Comment fixture cần bị xóa khỏi fixed file.", author="Fixture")
    doc.save(path)


def _build_tracked_changes(path: Path) -> None:
    doc = Document()
    _set_a4_margins(doc)
    doc.add_paragraph("Đoạn văn fixture có bật track revisions trong settings.xml.")
    doc.save(path)
    _inject_track_revisions_setting(path)


def _build_page_numbering_sections(path: Path) -> None:
    doc = Document()
    _set_a4_margins(doc)
    cover_section = doc.sections[0]
    _add_page_field(cover_section.footer.paragraphs[0])
    doc.add_paragraph("TRƯỜNG ĐẠI HỌC MẪU")
    doc.add_paragraph("ĐỒ ÁN TỐT NGHIỆP")

    front_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_a4_margins(front_section)
    _set_page_numbering(front_section, fmt="lowerRoman", start=1)
    _add_page_field(front_section.footer.paragraphs[0])
    doc.add_paragraph("LỜI CẢM ƠN")
    doc.add_paragraph(
        "Phần đầu tài liệu dùng số trang La Mã và không được làm lặp số i ở section sau."
    )

    repeated_roman_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_a4_margins(repeated_roman_section)
    _set_page_numbering(repeated_roman_section, fmt="lowerRoman", start=1)
    _add_page_field(repeated_roman_section.footer.paragraphs[0])
    doc.add_paragraph("LỜI CAM ĐOAN")
    doc.add_paragraph("Section này cố tình restart số La Mã để analyzer bắt lỗi.")

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_a4_margins(main_section)
    _set_page_numbering(main_section, fmt="lowerRoman", start=5)
    _add_page_field(main_section.footer.paragraphs[0])
    doc.add_paragraph("Chương 1.")
    doc.add_paragraph("TỔNG QUAN")
    doc.add_paragraph(
        "Đoạn nội dung chính đủ dài để kiểm tra section nội dung chính phải dùng số Ả Rập và bắt đầu từ 1."
    )
    doc.save(path)


def _build_advanced_manual_review(path: Path) -> None:
    doc = Document()
    _set_a4_margins(doc)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(
        "Đoạn này cố tình có màu chữ, highlight và gạch chân để rule nâng cao chỉ báo manual review."
    )
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.underline = True

    table = doc.add_table(rows=1, cols=1)
    cell_run = table.cell(0, 0).paragraphs[0].add_run("Ô bảng có highlight cần kiểm tra thủ công.")
    cell_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    equation_paragraph = doc.add_paragraph()
    equation_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    equation_paragraph.add_run("Công thức minh họa: ")
    equation_paragraph._p.append(
        parse_xml(
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            "<m:r><m:t>x+y=1</m:t></m:r>"
            "</m:oMath>"
        )
    )
    doc.save(path)


def _build_template_margin_variants(path: Path) -> None:
    doc = Document()
    _set_a4_margins(doc)
    doc.add_paragraph(
        "Fixture này dùng lề trên 2.5 cm, dưới 2.5 cm, trái 3.5 cm, phải 2.0 cm để kiểm tra margin theo template."
    )
    doc.save(path)


def _set_a4_margins(doc: Document) -> None:
    _set_section_a4_margins(doc.sections[0])


def _set_section_a4_margins(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2)


def _add_picture_to_paragraph(doc: Document, paragraph, *, width_cm: float) -> None:
    image_path = FIXTURE_DIR / "_tiny_fixture_image.png"
    image_path.write_bytes(TINY_PNG)
    try:
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
    finally:
        image_path.unlink(missing_ok=True)


def _add_page_field(paragraph) -> None:
    paragraph.paragraph_format.alignment = None
    run_begin = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_char_begin)

    run_instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    run_instr._r.append(instr_text)

    run_end = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_char_end)


def _set_page_numbering(section, *, fmt: str, start: int) -> None:
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    pg_num_type.set(qn("w:start"), str(start))


def _inject_track_revisions_setting(docx_path: Path) -> None:
    temp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as input_zip:
        with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as output_zip:
            for item in input_zip.infolist():
                content = input_zip.read(item.filename)
                if item.filename == "word/settings.xml" and b"<w:trackRevisions" not in content:
                    content = content.replace(b"</w:settings>", b"<w:trackRevisions/></w:settings>")
                output_zip.writestr(item, content)
    shutil.move(str(temp_path), docx_path)


if __name__ == "__main__":
    generated = ensure_docx_fixtures()
    for fixture_path in generated.values():
        print(fixture_path)
