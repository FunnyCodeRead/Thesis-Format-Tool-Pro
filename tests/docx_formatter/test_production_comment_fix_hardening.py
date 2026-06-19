from __future__ import annotations

import shutil
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Cm

from app.services.docx_formatter.analyzer import analyze_document_with_details
from app.services.docx_formatter.fixer import DocumentFixError, fix_document
from app.services.docx_formatter.rules.layout_abnormal_rule import LayoutAbnormalRule
from app.services.docx_formatter.utils.submission_cleanup import inspect_tracked_changes


class ProductionCommentFixHardeningTests(unittest.TestCase):
    def test_template_override_controls_margin_expectation(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            doc = Document()
            section = doc.sections[0]
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.5)
            section.right_margin = Cm(2.0)
            doc.add_paragraph("Doan van mau de test config margin.")
            doc.save(input_path)

            matching = analyze_document_with_details(
                str(input_path),
                "do_an_tot_nghiep",
                config_override={
                    "page_setup": {
                        "paper_size": "A4",
                        "margin_top_cm": 2.5,
                        "margin_bottom_cm": 2.5,
                        "margin_left_cm": 3.5,
                        "margin_right_cm": 2.0,
                    }
                },
            )
            strict = analyze_document_with_details(
                str(input_path),
                "do_an_tot_nghiep",
                config_override={
                    "page_setup": {
                        "paper_size": "A4",
                        "margin_top_cm": 2.0,
                        "margin_bottom_cm": 2.0,
                        "margin_left_cm": 3.5,
                        "margin_right_cm": 2.0,
                    }
                },
            )

            self.assertFalse(_has_margin_finding(matching["raw_findings"]))
            self.assertTrue(_has_margin_finding(strict["raw_findings"]))

    def test_layout_rule_splits_comments_from_track_changes(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "commented.docx"
            doc = Document()
            run = doc.add_paragraph("Doan co comment.").add_run(" Anchor")
            if not hasattr(doc, "add_comment"):
                self.skipTest("python-docx runtime does not support Document.add_comment")
            doc.add_comment([run], text="Comment test", author="Tester")
            doc.save(input_path)

            findings = LayoutAbnormalRule().analyze(Document(input_path), {"layout_abnormal": {"enabled": True}})
            finding_types = {finding.type for finding in findings}

            self.assertIn("COMMENTS_REVIEW", finding_types)
            self.assertNotIn("TRACK_CHANGES_REVIEW", finding_types)

    def test_fix_blocks_document_with_tracked_changes(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "tracked.docx"
            output_path = Path(temp_dir) / "fixed.docx"
            doc = Document()
            doc.add_paragraph("Doan van test track changes.")
            doc.save(input_path)
            _inject_track_revisions_setting(input_path)

            tracked = inspect_tracked_changes(str(input_path))
            self.assertGreater(tracked["total"], 0)

            with self.assertRaises(DocumentFixError):
                fix_document(
                    input_path=str(input_path),
                    document_type="do_an_tot_nghiep",
                    output_path=str(output_path),
                )


def _has_margin_finding(findings: list[dict]) -> bool:
    return any(finding.get("type") == "PAGE_MARGIN_ERROR" for finding in findings)


def _inject_track_revisions_setting(docx_path: Path) -> None:
    temp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as input_zip:
        with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as output_zip:
            for item in input_zip.infolist():
                content = input_zip.read(item.filename)
                if item.filename == "word/settings.xml" and b"<w:trackRevisions" not in content:
                    content = content.replace(b"</w:settings>", b"<w:trackRevisions/></w:settings>")
                output_zip.writestr(item, content)
    shutil.move(str(temp_path), docx_path)


if __name__ == "__main__":
    unittest.main()
