from __future__ import annotations

import unittest

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK
from docx.shared import Cm

from app.services.docx_formatter.rules.layout_abnormal_rule import LayoutAbnormalRule
from app.services.docx_formatter.rules.page_setup_rule import PageSetupRule
from app.services.docx_formatter.rules.region_review_rule import RegionReviewRule
from app.services.docx_formatter.rules.header_footer_format_rule import HeaderFooterFormatRule
from app.services.docx_formatter.rules.toc_structure_rule import TocStructureRule


class Phase2ProductionRuleTests(unittest.TestCase):
    def test_missing_toc_is_reported_when_document_has_chapter(self) -> None:
        doc = Document()
        doc.add_paragraph("Chuong 1.")
        doc.add_paragraph("TONG QUAN")

        findings = TocStructureRule().analyze(doc, _base_config())

        self.assertEqual([finding.type for finding in findings], ["TOC_MISSING_REVIEW"])
        self.assertFalse(findings[0].metadata["auto_fixable"])

    def test_typed_toc_without_field_is_manual_review(self) -> None:
        doc = Document()
        doc.add_paragraph("Muc luc")
        doc.add_paragraph("Chuong 1.")

        findings = TocStructureRule().analyze(doc, _base_config())

        self.assertEqual([finding.type for finding in findings], ["TOC_NOT_AUTOMATIC_REVIEW"])
        self.assertTrue(findings[0].metadata["manual_review"])

    def test_layout_abnormal_reports_manual_spacing_and_page_break(self) -> None:
        doc = Document()
        doc.add_paragraph("  Doan van can le bang dau cach thu cong trong phan noi dung.")
        page_break_paragraph = doc.add_paragraph("Doan co page break thu cong.")
        page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)

        findings = LayoutAbnormalRule().analyze(doc, _base_config())
        finding_types = {finding.type for finding in findings}

        self.assertIn("MANUAL_SPACING_REVIEW", finding_types)
        self.assertIn("MANUAL_PAGE_BREAK_REVIEW", finding_types)

    def test_landscape_section_is_manual_review_not_auto_fixable(self) -> None:
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)

        findings = PageSetupRule().analyze(doc, _base_config())
        landscape_findings = [
            finding for finding in findings if finding.type == "SECTION_LANDSCAPE_REVIEW"
        ]

        self.assertEqual(len(landscape_findings), 1)
        self.assertFalse(landscape_findings[0].metadata["auto_fixable"])

    def test_header_footer_font_is_checked_as_safe_format_issue(self) -> None:
        doc = Document()
        footer_paragraph = doc.sections[0].footer.paragraphs[0]
        run = footer_paragraph.add_run("Footer text")
        run.font.name = "Arial"

        findings = HeaderFooterFormatRule().analyze(doc, _base_config())
        font_findings = [
            finding for finding in findings if finding.type == "HEADER_FOOTER_FONT_NAME_ERROR"
        ]

        self.assertEqual(len(font_findings), 1)
        self.assertTrue(font_findings[0].metadata["auto_fixable"])


def _base_config() -> dict:
    return {
        "page_setup": {
            "paper_size": "A4",
            "margin_top_cm": 2.5,
            "margin_bottom_cm": 2.5,
            "margin_left_cm": 3.5,
            "margin_right_cm": 2.0,
        },
        "paragraph": {
            "font_name": "Times New Roman",
            "font_size": 13,
            "line_spacing": 1.3,
            "alignment": "JUSTIFY",
            "first_line_indent_cm": 1.0,
            "space_before_pt": 6,
            "space_after_pt": 6,
        },
        "toc": {"enabled": True},
        "layout_abnormal": {"enabled": True},
    }


if __name__ == "__main__":
    unittest.main()
