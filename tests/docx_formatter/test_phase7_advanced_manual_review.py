from __future__ import annotations

import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.shared import RGBColor

from app.services.docx_formatter.factories.rule_factory import RuleFactory
from app.services.docx_formatter.rules.advanced_format_rule import AdvancedFormatRule


class Phase7AdvancedManualReviewTests(unittest.TestCase):
    def test_text_decoration_is_manual_review_only(self) -> None:
        doc = Document()
        paragraph = doc.add_paragraph(
            "Doan van noi dung du dai de duoc kiem tra nhom mau chu highlight gach chan."
        )
        run = paragraph.runs[0]
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.font.underline = True

        findings = AdvancedFormatRule().analyze(doc, _config())

        self.assertEqual(len(findings), 1)
        self.assertEqual(findings[0].type, "TEXT_DECORATION_REVIEW")
        self.assertEqual(findings[0].metadata["report_group_id"], "text_decoration")
        self.assertFalse(findings[0].metadata["auto_fixable"])
        self.assertTrue(findings[0].metadata["manual_review"])
        self.assertEqual(findings[0].metadata["fix_action"]["type"], "manual_review")
        self.assertIn("highlight", findings[0].current_value or "")

    def test_table_text_decoration_is_reported_at_table_cell_locator(self) -> None:
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        run = table.cell(0, 0).paragraphs[0].add_run("Chu trong bang bi highlight.")
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        findings = AdvancedFormatRule().analyze(doc, _config())

        self.assertEqual(len(findings), 1)
        self.assertEqual(findings[0].type, "TABLE_TEXT_DECORATION_REVIEW")
        self.assertEqual(findings[0].metadata["target"], "table_cell")
        self.assertEqual(findings[0].metadata["table_index"], 1)
        self.assertEqual(findings[0].metadata["row_index"], 1)
        self.assertEqual(findings[0].metadata["cell_index"], 1)

    def test_equation_layout_is_manual_review(self) -> None:
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph._p.append(
            parse_xml(
                '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                "<m:r><m:t>x+y=1</m:t></m:r>"
                "</m:oMath>"
            )
        )

        findings = AdvancedFormatRule().analyze(doc, _config())

        self.assertEqual(len(findings), 1)
        self.assertEqual(findings[0].type, "EQUATION_LAYOUT_REVIEW")
        self.assertEqual(findings[0].metadata["report_group_id"], "equation_layout")
        self.assertFalse(findings[0].metadata["auto_fixable"])
        self.assertTrue(findings[0].metadata["manual_review"])

    def test_advanced_rule_can_be_disabled_by_config(self) -> None:
        doc = Document()
        run = doc.add_paragraph("Doan van co highlight.").runs[0]
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        findings = AdvancedFormatRule().analyze(doc, {"advanced_review": {"enabled": False}})

        self.assertEqual(findings, [])

    def test_fix_registry_does_not_include_advanced_rule(self) -> None:
        analyze_rule_types = {type(rule) for rule in RuleFactory.create_analyze_rules()}
        fix_rule_types = {type(rule) for rule in RuleFactory.create_fix_rules()}

        self.assertIn(AdvancedFormatRule, analyze_rule_types)
        self.assertNotIn(AdvancedFormatRule, fix_rule_types)


def _config() -> dict:
    return {"advanced_review": {"enabled": True}}


if __name__ == "__main__":
    unittest.main()
