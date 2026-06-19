from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.services.docx_formatter.analyzer import analyze_document_with_details
from app.services.docx_formatter.factories.rule_factory import RuleFactory
from app.services.docx_formatter.fixer import fix_document
from app.services.docx_formatter.rules.caption_numbering_rule import CaptionNumberingRule
from app.services.docx_formatter.rules.page_numbering_rule import PageNumberingRule
from app.services.docx_formatter.rules.region_review_rule import RegionReviewRule


class Phase3SafeFixExpandedTests(unittest.TestCase):
    def test_front_matter_table_cell_and_header_footer_are_safe_fixable(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "fixed.docx"

            doc = Document()
            front_matter = doc.add_paragraph("LỜI CẢM ƠN")
            front_matter.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            front_run = front_matter.runs[0]
            front_run.font.name = "Arial"
            front_run.font.size = Pt(12)
            front_run.font.bold = False

            body = doc.add_paragraph(
                "Đây là đoạn văn nội dung đủ dài để hệ thống nhận diện body paragraph và kiểm tra định dạng."
            )
            body.add_run(" Nội dung không được thay đổi.")

            table = doc.add_table(rows=1, cols=1)
            table_run = table.cell(0, 0).paragraphs[0].add_run("Nội dung ô bảng giữ nguyên.")
            table_run.font.name = "Arial"
            table_run.font.size = Pt(11)

            footer_run = doc.sections[0].footer.paragraphs[0].add_run("5")
            footer_run.font.name = "Arial"
            footer_run.font.size = Pt(11)
            doc.save(input_path)

            original_text = _visible_text(Document(input_path))
            analysis_before = analyze_document_with_details(str(input_path), "do_an_tot_nghiep")
            before_types = {finding["type"] for finding in analysis_before["raw_findings"]}

            result = fix_document(
                input_path=str(input_path),
                document_type="do_an_tot_nghiep",
                output_path=str(output_path),
            )
            fixed_doc = Document(output_path)
            analysis_after = analyze_document_with_details(str(output_path), "do_an_tot_nghiep")
            after_types = {finding["type"] for finding in analysis_after["raw_findings"]}

            self.assertIn("FRONT_MATTER_HEADING_ALIGNMENT_ERROR", before_types)
            self.assertIn("TABLE_CELL_FONT_NAME_ERROR", before_types)
            self.assertIn("HEADER_FOOTER_FONT_NAME_ERROR", before_types)

            self.assertNotIn("FRONT_MATTER_HEADING_ALIGNMENT_ERROR", after_types)
            self.assertNotIn("TABLE_CELL_FONT_NAME_ERROR", after_types)
            self.assertNotIn("TABLE_CELL_FONT_SIZE_ERROR", after_types)
            self.assertNotIn("HEADER_FOOTER_FONT_NAME_ERROR", after_types)
            self.assertNotIn("HEADER_FOOTER_FONT_SIZE_ERROR", after_types)

            self.assertEqual(_visible_text(fixed_doc), original_text)
            self.assertGreater(result["front_matter_heading_changes"], 0)
            self.assertGreater(result["table_cell_changes"], 0)
            self.assertGreater(result["header_footer_changes"], 0)
            self.assertTrue(result["safety_checks"]["visible_text_preserved"])
            self.assertEqual(fixed_doc.paragraphs[0].paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(fixed_doc.tables[0].cell(0, 0).paragraphs[0].runs[0].font.name, "Times New Roman")
            self.assertEqual(fixed_doc.sections[0].footer.paragraphs[0].runs[0].font.name, "Times New Roman")

    def test_structural_review_rules_remain_blocked_from_fix_registry(self) -> None:
        fix_rule_types = {type(rule) for rule in RuleFactory.create_fix_rules()}

        self.assertNotIn(CaptionNumberingRule, fix_rule_types)
        self.assertNotIn(PageNumberingRule, fix_rule_types)
        self.assertNotIn(RegionReviewRule, fix_rule_types)


def _visible_text(doc: Document) -> dict:
    return {
        "body_paragraphs": [paragraph.text for paragraph in doc.paragraphs],
        "tables": [
            [
                [
                    [paragraph.text for paragraph in cell.paragraphs]
                    for cell in row.cells
                ]
                for row in table.rows
            ]
            for table in doc.tables
        ],
        "headers": [
            [paragraph.text for paragraph in section.header.paragraphs]
            for section in doc.sections
        ],
        "footers": [
            [paragraph.text for paragraph in section.footer.paragraphs]
            for section in doc.sections
        ],
    }


if __name__ == "__main__":
    unittest.main()
