from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from app.services.docx_formatter.fixer import DocumentFixError, fix_document
from app.services.docx_formatter.utils.docx_units import length_to_cm


class SelectiveFixScopeTests(unittest.TestCase):
    def test_safe_scope_page_setup_only_does_not_fix_paragraph_format(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "fixed.docx"

            doc = _build_document_with_bad_margin_and_body()
            doc.save(input_path)

            result = fix_document(
                input_path=str(input_path),
                document_type="do_an_tot_nghiep",
                output_path=str(output_path),
                fix_options={
                    "fix_mode": "safe_scope",
                    "fix_scope": ["page_setup"],
                },
            )

            fixed_doc = Document(output_path)
            self.assertEqual(result["fix_mode"], "safe_scope")
            self.assertEqual(result["applied_fix_scope"], ["page_setup"])
            self.assertEqual(result["safe_fix_rules"], ["PageSetupRule"])
            self.assertIn("ParagraphFormatRule", result["skipped_safe_fix_rules"])
            self.assertGreater(result["page_setup_changes"], 0)
            self.assertEqual(result["paragraph_changes"], 0)
            self.assertEqual(result["post_fix_validation"]["status"], "passed")
            self.assertEqual(
                result["post_fix_validation"]["remaining_selected_safe_issue_count"],
                0,
            )
            self.assertGreater(
                result["post_fix_validation"]["remaining_unselected_safe_issue_count"],
                0,
            )
            self.assertAlmostEqual(length_to_cm(fixed_doc.sections[0].left_margin), 3.5, places=1)
            self.assertEqual(
                fixed_doc.paragraphs[0].paragraph_format.alignment,
                WD_ALIGN_PARAGRAPH.LEFT,
            )

    def test_safe_scope_paragraph_only_does_not_fix_page_setup(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "fixed.docx"

            doc = _build_document_with_bad_margin_and_body()
            doc.save(input_path)

            result = fix_document(
                input_path=str(input_path),
                document_type="do_an_tot_nghiep",
                output_path=str(output_path),
                fix_options={
                    "fix_mode": "safe_scope",
                    "fix_scope": ["paragraph_format"],
                },
            )

            fixed_doc = Document(output_path)
            self.assertEqual(result["applied_fix_scope"], ["paragraph_format"])
            self.assertEqual(result["safe_fix_rules"], ["ParagraphFormatRule"])
            self.assertIn("PageSetupRule", result["skipped_safe_fix_rules"])
            self.assertEqual(result["page_setup_changes"], 0)
            self.assertGreater(result["paragraph_changes"] + result["style_changes"], 0)
            self.assertEqual(result["post_fix_validation"]["status"], "passed")
            self.assertEqual(
                result["post_fix_validation"]["remaining_selected_safe_issue_count"],
                0,
            )
            self.assertGreater(
                result["post_fix_validation"]["remaining_unselected_safe_issue_count"],
                0,
            )
            self.assertAlmostEqual(length_to_cm(fixed_doc.sections[0].left_margin), 1.0, places=1)
            self.assertEqual(
                fixed_doc.paragraphs[0].paragraph_format.alignment,
                WD_ALIGN_PARAGRAPH.JUSTIFY,
            )

    def test_invalid_safe_scope_is_rejected_before_writing_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "fixed.docx"

            doc = _build_document_with_bad_margin_and_body()
            doc.save(input_path)

            with self.assertRaises(DocumentFixError):
                fix_document(
                    input_path=str(input_path),
                    document_type="do_an_tot_nghiep",
                    output_path=str(output_path),
                    fix_options={
                        "fix_mode": "safe_scope",
                        "fix_scope": ["caption_numbering"],
                    },
                )

            self.assertFalse(output_path.exists())

    def test_post_fix_gate_rejects_remaining_selected_safe_issues(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "fixed.docx"

            doc = _build_document_with_bad_margin_and_body()
            doc.save(input_path)

            post_fix_analysis = {
                "raw_findings": [
                    {
                        "type": "PARAGRAPH_ALIGNMENT_ERROR",
                        "severity": "warning",
                        "location": "Paragraph 1",
                        "message": "Căn lề đoạn văn chưa đúng yêu cầu.",
                        "metadata": {
                            "fixability_scope": "safe_auto_fix",
                            "report_group_id": "body_paragraph",
                        },
                    }
                ],
                "render_verification": {"status": "skipped"},
            }

            with patch(
                "app.services.docx_formatter.fixer.analyze_document_with_details",
                return_value=post_fix_analysis,
            ):
                with self.assertRaises(DocumentFixError) as caught:
                    fix_document(
                        input_path=str(input_path),
                        document_type="do_an_tot_nghiep",
                        output_path=str(output_path),
                        fix_options={
                            "fix_mode": "safe_scope",
                            "fix_scope": ["paragraph_format"],
                        },
                    )

            self.assertIn("Post-fix analyzer gate failed", str(caught.exception))


def _build_document_with_bad_margin_and_body() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    paragraph = doc.add_paragraph(
        "Day la doan van noi dung du dai de duoc xem la body paragraph "
        "va can duoc dinh dang lai theo chuan cua template."
    )
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return doc


if __name__ == "__main__":
    unittest.main()
