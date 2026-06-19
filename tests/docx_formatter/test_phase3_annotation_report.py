from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from app.services.docx_formatter.engine.annotation_engine import AnnotationEngine
from app.services.docx_formatter.writers.word_comment_writer import WordCommentWriter


class Phase3AnnotationReportTests(unittest.TestCase):
    def test_engine_groups_same_paragraph_even_when_raw_location_differs(self) -> None:
        findings = [
            _finding("PARAGRAPH_ALIGNMENT_ERROR", paragraph_index=2, field="alignment", location="Paragraph 2"),
            _finding("PARAGRAPH_FONT_SIZE_ERROR", paragraph_index=2, field="font_size", location="Body paragraph 2"),
        ]

        comments = AnnotationEngine().build_comments(findings)

        self.assertEqual(len(comments), 1)
        self.assertEqual(comments[0].source_count, 2)

    def test_engine_uses_hybrid_style_comment_for_repeated_format_findings(self) -> None:
        findings = [
            _auto_fixable_finding(
                "PARAGRAPH_ALIGNMENT_ERROR",
                paragraph_index=index,
                field="alignment",
                style_name="Normal",
            )
            for index in range(1, 51)
        ]

        comments = AnnotationEngine().build_comments(findings)

        self.assertLess(len(comments), 10)
        self.assertEqual(sum(comment.source_count for comment in comments), 50)
        self.assertEqual(comments[0].source_count, 50)
        self.assertIn("Normal", comments[0].to_text())
        self.assertIn("kiểu định dạng", comments[0].to_text().lower())

    def test_writer_creates_comment_for_each_error_paragraph_across_document(self) -> None:
        if not hasattr(Document(), "add_comment"):
            self.skipTest("python-docx runtime does not support Document.add_comment")

        findings = [
            _finding("PARAGRAPH_ALIGNMENT_ERROR", paragraph_index=1, field="alignment"),
            _finding("PARAGRAPH_FONT_SIZE_ERROR", paragraph_index=2, field="font_size"),
            _finding("PARAGRAPH_LINE_SPACING_ERROR", paragraph_index=3, field="line_spacing"),
        ]
        comments = AnnotationEngine().build_comments(findings)

        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "annotated.docx"
            doc = Document()
            doc.add_paragraph("Doan van thu nhat co loi.")
            doc.add_paragraph("Doan van thu hai co loi.")
            doc.add_paragraph("Doan van thu ba co loi.")
            doc.save(input_path)

            result = WordCommentWriter().write(
                input_path=str(input_path),
                output_path=str(output_path),
                comments=comments,
            )

            self.assertEqual(result["total_findings"], 3)
            self.assertEqual(result["total_comments_created"], 3)
            self.assertEqual(result["skipped_findings"], 0)

    def test_writer_reports_skipped_findings_with_reason_and_ids(self) -> None:
        if not hasattr(Document(), "add_comment"):
            self.skipTest("python-docx runtime does not support Document.add_comment")

        findings = [_finding("PARAGRAPH_ALIGNMENT_ERROR", paragraph_index=99, field="alignment")]
        comments = AnnotationEngine().build_comments(findings)

        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "annotated.docx"
            doc = Document()
            doc.add_paragraph("Doan duy nhat trong file.")
            doc.save(input_path)

            result = WordCommentWriter().write(
                input_path=str(input_path),
                output_path=str(output_path),
                comments=comments,
            )

            self.assertEqual(result["total_findings"], 1)
            self.assertEqual(result["total_comments_created"], 0)
            self.assertEqual(result["skipped_findings"], 1)
            self.assertEqual(result["skipped_comments"], 1)
            self.assertEqual(result["skipped_reason"][0]["finding_ids"], ["FINDING-0001"])
            self.assertEqual(result["skipped_reason"][0]["source_types"], ["PARAGRAPH_ALIGNMENT_ERROR"])

    def test_writer_anchors_empty_paragraph_finding_to_nearby_commentable_paragraph(self) -> None:
        if not hasattr(Document(), "add_comment"):
            self.skipTest("python-docx runtime does not support Document.add_comment")

        findings = [_finding("EXCESSIVE_BLANK_PARAGRAPHS_REVIEW", paragraph_index=2, field="blank_paragraphs")]
        comments = AnnotationEngine().build_comments(findings)

        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "annotated.docx"
            doc = Document()
            doc.add_paragraph("Doan truoc do.")
            doc.add_paragraph("")
            doc.add_paragraph("Doan sau do.")
            doc.save(input_path)

            result = WordCommentWriter().write(
                input_path=str(input_path),
                output_path=str(output_path),
                comments=comments,
            )

            self.assertEqual(result["total_findings"], 1)
            self.assertEqual(result["total_comments_created"], 1)
            self.assertEqual(result["skipped_findings"], 0)

    def test_table_cell_target_falls_back_to_first_commentable_table_cell(self) -> None:
        if not hasattr(Document(), "add_comment"):
            self.skipTest("python-docx runtime does not support Document.add_comment")

        findings = [
            _finding(
                "TABLE_WIDTH_REVIEW",
                target="table_cell",
                field="table_width",
                table_index=1,
                row_index=1,
                cell_index=1,
                table_paragraph_index=1,
            )
        ]
        comments = AnnotationEngine().build_comments(findings)

        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "annotated.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 1).paragraphs[0].add_run("O thu hai co noi dung.")
            doc.save(input_path)

            result = WordCommentWriter().write(
                input_path=str(input_path),
                output_path=str(output_path),
                comments=comments,
            )

            self.assertEqual(result["total_findings"], 1)
            self.assertEqual(result["total_comments_created"], 1)
            self.assertEqual(result["skipped_findings"], 0)

    def test_footer_target_uses_part_paragraph_index(self) -> None:
        if not hasattr(Document(), "add_comment"):
            self.skipTest("python-docx runtime does not support Document.add_comment")

        findings = [
            _finding(
                "HEADER_FOOTER_FONT_NAME_ERROR",
                target="footer",
                field="font_name",
                section_index=1,
                part_paragraph_index=2,
            )
        ]
        comments = AnnotationEngine().build_comments(findings)

        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "annotated.docx"
            doc = Document()
            doc.add_paragraph("Body")
            footer = doc.sections[0].footer
            footer.paragraphs[0].add_run("Chan trang thu nhat.")
            footer.add_paragraph("Chan trang thu hai.")
            doc.save(input_path)

            result = WordCommentWriter().write(
                input_path=str(input_path),
                output_path=str(output_path),
                comments=comments,
            )

            self.assertEqual(result["total_comments_created"], 1)
            with zipfile.ZipFile(output_path, "r") as docx_zip:
                footer_xml = docx_zip.read("word/footer1.xml").decode("utf-8")

            self.assertIn("commentRangeStart", footer_xml)
            self.assertIn("Chan trang thu hai.", footer_xml)


def _finding(
    finding_type: str,
    *,
    paragraph_index: int | None = None,
    field: str,
    location: str | None = None,
    target: str = "paragraph",
    section_index: int | None = None,
    table_index: int | None = None,
    row_index: int | None = None,
    cell_index: int | None = None,
    table_paragraph_index: int | None = None,
    part_paragraph_index: int | None = None,
) -> dict:
    metadata = {
        "target": target,
        "context": target,
        "report_group_id": "body_paragraph" if target == "paragraph" else target,
        "field": field,
        "auto_fixable": False,
        "manual_review": True,
        "fix_action": {"type": "manual_review"},
    }
    optional_values = {
        "paragraph_index": paragraph_index,
        "section_index": section_index,
        "table_index": table_index,
        "row_index": row_index,
        "cell_index": cell_index,
        "table_paragraph_index": table_paragraph_index,
        "part_paragraph_index": part_paragraph_index,
    }
    metadata.update({key: value for key, value in optional_values.items() if value is not None})

    return {
        "type": finding_type,
        "severity": "warning",
        "location": location or (f"Paragraph {paragraph_index}" if paragraph_index else "Document"),
        "message": "Formatting issue.",
        "current_value": "Sai",
        "expected_value": "Dung",
        "suggestion": "Kiem tra dinh dang.",
        "metadata": metadata,
    }


def _auto_fixable_finding(
    finding_type: str,
    *,
    paragraph_index: int,
    field: str,
    style_name: str,
) -> dict:
    metadata = {
        "target": "paragraph",
        "context": "body_paragraph",
        "report_group_id": "body_paragraph",
        "paragraph_index": paragraph_index,
        "field": field,
        "style_name": style_name,
        "auto_fixable": True,
        "manual_review": False,
        "fix_action": {"type": "paragraph_format"},
    }
    return {
        "type": finding_type,
        "severity": "warning",
        "location": f"Paragraph {paragraph_index}",
        "message": "Paragraph alignment does not match the required format.",
        "current_value": "LEFT",
        "expected_value": "JUSTIFY",
        "suggestion": "Set alignment to JUSTIFY.",
        "metadata": metadata,
    }


if __name__ == "__main__":
    unittest.main()
