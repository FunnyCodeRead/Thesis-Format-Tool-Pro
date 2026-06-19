from __future__ import annotations

import copy
import json
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from app.services.docx_formatter.annotator import annotate_document
from app.services.docx_formatter.analyzer import analyze_document_with_details
from app.services.docx_formatter.engine.report_builder import ReportBuilder
from app.services.docx_formatter.fixer import DocumentFixError, fix_document
from app.services.docx_formatter.utils.submission_cleanup import (
    inspect_submission_artifacts,
    inspect_tracked_changes,
)
from tests.fixtures.docx_fixtures import ensure_docx_fixtures


class Phase3DocxFixturePackTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.fixtures = ensure_docx_fixtures()

    def test_fixture_pack_materializes_expected_docx_files(self) -> None:
        expected = {
            "style_inheritance_ok",
            "cover_image_without_caption",
            "body_image_without_caption",
            "hand_toc_bad_caption_numbers",
            "commented_highlighted",
            "tracked_changes",
            "page_numbering_sections",
            "advanced_manual_review",
            "template_margin_variants",
        }

        self.assertEqual(set(self.fixtures), expected)
        for path in self.fixtures.values():
            self.assertTrue(path.exists(), path)
            self.assertGreater(path.stat().st_size, 0)
            Document(path)

    def test_style_inherited_font_size_does_not_create_body_format_false_positive(self) -> None:
        result = _analyze_fixture(self.fixtures["style_inheritance_ok"])
        finding_types = _finding_types(result)

        self.assertNotIn("PARAGRAPH_FONT_SIZE_ERROR", finding_types)
        self.assertNotIn("PARAGRAPH_FONT_NAME_ERROR", finding_types)
        self.assertNotIn("PARAGRAPH_LINE_SPACING_ERROR", finding_types)
        self.assertNotIn("PARAGRAPH_SPACE_BEFORE_PT_ERROR", finding_types)
        self.assertNotIn("PARAGRAPH_SPACE_AFTER_PT_ERROR", finding_types)

    def test_cover_logo_image_is_not_reported_as_missing_caption(self) -> None:
        result = _analyze_fixture(self.fixtures["cover_image_without_caption"])

        self.assertNotIn("IMAGE_LAYOUT_REVIEW", _finding_types(result))

    def test_body_image_without_nearby_caption_is_reported(self) -> None:
        result = _analyze_fixture(self.fixtures["body_image_without_caption"])
        image_findings = [
            finding
            for finding in result["raw_findings"]
            if finding.get("type") == "IMAGE_LAYOUT_REVIEW"
        ]

        self.assertGreaterEqual(len(image_findings), 1)
        self.assertTrue(
            all(finding.get("metadata", {}).get("paragraph_index") for finding in image_findings)
        )

    def test_hand_toc_and_caption_numbering_fixture_reports_structure_issues(self) -> None:
        result = _analyze_fixture(self.fixtures["hand_toc_bad_caption_numbers"])
        finding_types = _finding_types(result)

        self.assertIn("TOC_NOT_AUTOMATIC_REVIEW", finding_types)
        self.assertIn("FIGURE_NUMBERING_MALFORMED_REVIEW", finding_types)
        self.assertIn("TABLE_LIST_DUPLICATE_REVIEW", finding_types)

    def test_commented_highlighted_fixture_fixes_to_clean_submission_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "fixed.docx"
            original_text = _visible_text(Document(self.fixtures["commented_highlighted"]))

            result = fix_document(
                input_path=str(self.fixtures["commented_highlighted"]),
                document_type="do_an_tot_nghiep",
                output_path=str(output_path),
                config_override=_fixture_config(),
            )

            self.assertTrue(result["safety_checks"]["visible_text_preserved"])
            self.assertTrue(result["safety_checks"]["comments_removed"])
            self.assertTrue(result["safety_checks"]["highlights_removed"])
            self.assertEqual(_visible_text(Document(output_path)), original_text)

            artifacts = inspect_submission_artifacts(str(output_path))
            self.assertEqual(artifacts["comment_parts"], 0)
            self.assertEqual(artifacts["comment_markers"], 0)
            self.assertEqual(artifacts["highlights"], 0)
            with zipfile.ZipFile(output_path, "r") as docx_zip:
                self.assertNotIn("word/comments.xml", set(docx_zip.namelist()))

    def test_tracked_changes_fixture_blocks_safe_fix(self) -> None:
        tracked_report = inspect_tracked_changes(str(self.fixtures["tracked_changes"]))
        self.assertGreater(tracked_report["total"], 0)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "fixed.docx"
            with self.assertRaises(DocumentFixError):
                fix_document(
                    input_path=str(self.fixtures["tracked_changes"]),
                    document_type="do_an_tot_nghiep",
                    output_path=str(output_path),
                    config_override=_fixture_config(),
                )

    def test_page_numbering_sections_fixture_reports_section_numbering_issues(self) -> None:
        result = _analyze_fixture(self.fixtures["page_numbering_sections"])
        finding_types = _finding_types(result)

        self.assertIn("COVER_PAGE_NUMBER_VISIBLE_REVIEW", finding_types)
        self.assertIn("ROMAN_PAGE_NUMBER_REPEATED_REVIEW", finding_types)
        self.assertIn("MAIN_PAGE_NUMBER_FORMAT_REVIEW", finding_types)
        self.assertIn("MAIN_PAGE_NUMBER_RESET_REVIEW", finding_types)

        page_number_findings = [
            finding
            for finding in result["raw_findings"]
            if finding.get("type")
            in {
                "COVER_PAGE_NUMBER_VISIBLE_REVIEW",
                "ROMAN_PAGE_NUMBER_REPEATED_REVIEW",
                "MAIN_PAGE_NUMBER_FORMAT_REVIEW",
                "MAIN_PAGE_NUMBER_RESET_REVIEW",
            }
        ]
        self.assertTrue(page_number_findings)
        for finding in page_number_findings:
            metadata = finding.get("metadata", {})
            self.assertTrue(metadata.get("manual_review"), finding)
            self.assertFalse(metadata.get("auto_fixable"), finding)
            self.assertEqual(metadata.get("fix_action", {}).get("type"), "manual_review")
            self.assertIsNotNone(metadata.get("section_index"), finding)

    def test_advanced_manual_review_fixture_reports_decoration_table_and_equation(self) -> None:
        result = _analyze_fixture(self.fixtures["advanced_manual_review"])
        finding_types = _finding_types(result)

        self.assertIn("TEXT_DECORATION_REVIEW", finding_types)
        self.assertIn("TABLE_TEXT_DECORATION_REVIEW", finding_types)
        self.assertIn("EQUATION_LAYOUT_REVIEW", finding_types)

        advanced_findings = [
            finding
            for finding in result["raw_findings"]
            if finding.get("type")
            in {
                "TEXT_DECORATION_REVIEW",
                "TABLE_TEXT_DECORATION_REVIEW",
                "EQUATION_LAYOUT_REVIEW",
            }
        ]
        self.assertTrue(advanced_findings)
        for finding in advanced_findings:
            metadata = finding.get("metadata", {})
            self.assertTrue(metadata.get("manual_review"), finding)
            self.assertFalse(metadata.get("auto_fixable"), finding)
            self.assertEqual(metadata.get("fix_action", {}).get("type"), "manual_review")

    def test_template_margin_fixture_passes_and_fails_by_selected_template(self) -> None:
        matching_result = _analyze_fixture(self.fixtures["template_margin_variants"])
        self.assertNotIn("PAGE_MARGIN_ERROR", _finding_types(matching_result))

        strict_config = copy.deepcopy(_fixture_config())
        strict_config["page_setup"]["margin_top_cm"] = 2.0
        strict_config["page_setup"]["margin_bottom_cm"] = 2.0
        strict_result = analyze_document_with_details(
            str(self.fixtures["template_margin_variants"]),
            "do_an_tot_nghiep",
            config_override=strict_config,
        )
        strict_margin_findings = [
            finding
            for finding in strict_result["raw_findings"]
            if finding.get("type") == "PAGE_MARGIN_ERROR"
        ]
        self.assertGreaterEqual(len(strict_margin_findings), 2)
        self.assertTrue(
            any(
                finding.get("metadata", {}).get("field") == "margin_top_cm"
                for finding in strict_margin_findings
            )
        )
        self.assertTrue(
            any(
                finding.get("metadata", {}).get("field") == "margin_bottom_cm"
                for finding in strict_margin_findings
            )
        )

    def test_annotated_report_fixture_creates_real_word_comments_without_anchor_skips(self) -> None:
        if not hasattr(Document(), "add_comment"):
            self.skipTest("python-docx runtime does not support Document.add_comment")

        result = _analyze_fixture(self.fixtures["hand_toc_bad_caption_numbers"])
        anchorable_findings = [
            finding
            for finding in result["raw_findings"]
            if finding.get("metadata", {}).get("paragraph_index")
        ]
        self.assertGreater(len(anchorable_findings), 0)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "annotated.docx"
            annotate_result = annotate_document(
                input_path=str(self.fixtures["hand_toc_bad_caption_numbers"]),
                output_path=str(output_path),
                findings=anchorable_findings,
            )

            self.assertEqual(annotate_result["comment_strategy"], "hybrid")
            self.assertEqual(annotate_result["skipped_comments"], 0)
            self.assertEqual(annotate_result["skipped_findings"], 0)
            self.assertGreater(annotate_result["total_findings"], 0)
            self.assertGreater(annotate_result["total_comments_created"], 0)
            self.assertTrue(output_path.exists())

            with zipfile.ZipFile(output_path, "r") as docx_zip:
                names = set(docx_zip.namelist())
                document_xml = docx_zip.read("word/document.xml").decode("utf-8")
                self.assertIn("word/comments.xml", names)
                self.assertIn("w:commentRangeStart", document_xml)
                self.assertIn("w:commentRangeEnd", document_xml)
                self.assertIn("w:commentReference", document_xml)

    def test_annotated_report_anchors_table_equation_and_page_number_findings(self) -> None:
        if not hasattr(Document(), "add_comment"):
            self.skipTest("python-docx runtime does not support Document.add_comment")

        scenarios = [
            (
                "advanced_manual_review",
                {
                    "TEXT_DECORATION_REVIEW",
                    "TABLE_TEXT_DECORATION_REVIEW",
                    "EQUATION_LAYOUT_REVIEW",
                },
            ),
            (
                "page_numbering_sections",
                {
                    "COVER_PAGE_NUMBER_VISIBLE_REVIEW",
                    "PAGE_NUMBER_ALIGNMENT_REVIEW",
                    "ROMAN_PAGE_NUMBER_REPEATED_REVIEW",
                    "MAIN_PAGE_NUMBER_FORMAT_REVIEW",
                    "MAIN_PAGE_NUMBER_RESET_REVIEW",
                },
            ),
        ]

        with tempfile.TemporaryDirectory() as temp_dir:
            for fixture_name, finding_types in scenarios:
                result = _analyze_fixture(self.fixtures[fixture_name])
                selected_findings = [
                    finding
                    for finding in result["raw_findings"]
                    if finding.get("type") in finding_types
                ]
                self.assertGreater(len(selected_findings), 0, fixture_name)

                output_path = Path(temp_dir) / f"{fixture_name}-annotated.docx"
                annotate_result = annotate_document(
                    input_path=str(self.fixtures[fixture_name]),
                    output_path=str(output_path),
                    findings=selected_findings,
                )

                self.assertEqual(
                    annotate_result["total_findings"],
                    len(selected_findings),
                    fixture_name,
                )
                self.assertEqual(annotate_result["skipped_comments"], 0, fixture_name)
                self.assertEqual(annotate_result["skipped_findings"], 0, fixture_name)
                self.assertGreater(
                    annotate_result["total_comments_created"],
                    0,
                    fixture_name,
                )
                self.assertEqual(
                    annotate_result["total_findings"],
                    annotate_result["total_comments_created"]
                    + annotate_result["grouped_findings_in_comments"],
                    fixture_name,
                )

    def test_production_report_contract_from_fixture_is_grouped_and_user_facing(self) -> None:
        result = _analyze_fixture(self.fixtures["hand_toc_bad_caption_numbers"])
        report = ReportBuilder().build(
            raw_findings=result["raw_findings"],
            document_id="fixture-document",
            document_type="do_an_tot_nghiep",
            filename="hand_toc_bad_caption_numbers.docx",
            template_name="Fixture template",
        )

        self.assertTrue(report["ok"])
        self.assertEqual(report["reference"]["mode"], "style_only")
        self.assertIn("style_fix_groups", report["summary"])
        self.assertIn("issue_groups", report)
        self.assertGreater(len(report["issue_groups"]), 0)
        self.assertNotIn("examples", json.dumps(report, ensure_ascii=False))

        for group in report["issue_groups"]:
            self.assertIn(group["recommended_fix_scope"], {"style", "paragraph", "manual_review"})
            self.assertIsInstance(group["affected_styles"], list)
            self.assertGreater(len(group["issues"]), 0)
            for issue in group["issues"]:
                self.assertIn("issue_id", issue)
                self.assertIn("fix_action", issue)
                self.assertIn("location", issue)
                self.assertIn("target", issue)
                self.assertIn("rule", issue)
                self.assertIn(issue["severity"], {"critical", "major", "minor"})

        user_facing_text = _collect_user_facing_report_text(report)
        _assert_no_banned_user_facing_text(user_facing_text)

    def test_report_builder_normalizes_legacy_no_accent_finding_text(self) -> None:
        report = ReportBuilder().build(
            raw_findings=[
                {
                    "type": "EXCESSIVE_BLANK_PARAGRAPHS_REVIEW",
                    "severity": "warning",
                    "location": "Paragraph 3",
                    "message": "Co nhieu dong trong lien tiep lam bo cuc co the bi vo.",
                    "current_value": "Dong trong lien tiep: 4 dong trong lien tiep",
                    "expected_value": "Dong trong lien tiep: Khong qua 2 dong trong lien tiep trong phan noi dung.",
                    "suggestion": "Kiem tra cac dong trong thu cong; neu can xuong trang thi dung page break hoac section dung cach.",
                    "metadata": {
                        "target": "paragraph",
                        "context": "layout_abnormal",
                        "report_group_id": "layout_abnormal",
                        "report_severity": "minor",
                        "paragraph_index": 3,
                        "field": "blank_paragraphs",
                        "manual_review": True,
                        "auto_fixable": False,
                        "fix_action": {"type": "manual_review", "reason": "manual"},
                    },
                }
            ],
            document_id="fixture-document",
            document_type="do_an_tot_nghiep",
            filename="legacy.docx",
        )
        text = _collect_user_facing_report_text(report)

        self.assertIn("Dòng trống liên tiếp", text)
        self.assertIn("Có nhiều dòng trống liên tiếp làm bố cục có thể bị vỡ.", text)
        self.assertIn("Không quá 2 dòng trống liên tiếp trong phần nội dung.", text)
        _assert_no_banned_user_facing_text(text)


def _analyze_fixture(path: Path) -> dict:
    return analyze_document_with_details(
        str(path),
        "do_an_tot_nghiep",
        config_override=_fixture_config(),
    )


def _fixture_config() -> dict:
    return {
        "page_setup": {
            "paper_size": "A4",
            "margin_top_cm": 2.5,
            "margin_bottom_cm": 2.5,
            "margin_left_cm": 3.5,
            "margin_right_cm": 2.0,
        },
        "paragraph": {
            "font_name": "Times New Roman",
            "font_size": 13,
            "line_spacing": 1.3,
            "alignment": "JUSTIFY",
            "first_line_indent_cm": 1.0,
            "space_before_pt": 6,
            "space_after_pt": 6,
        },
        "list_item": {
            "font_name": "Times New Roman",
            "font_size": 13,
            "line_spacing": 1.3,
            "space_before_pt": 6,
            "space_after_pt": 6,
        },
        "front_matter_heading": {
            "font_name": "Times New Roman",
            "font_size": 13,
            "bold": True,
            "uppercase": True,
            "alignment": "CENTER",
            "space_after_pt": 6,
        },
        "chapter_layout": {"enabled": True},
        "pagination": {"enabled": True},
        "caption_numbering": {
            "enabled": True,
            "min_number_parts": 2,
            "max_number_parts": 2,
            "require_chapter_prefix": True,
        },
        "toc": {"enabled": True},
        "layout_abnormal": {"enabled": True},
        "render_verification": {"enabled": False},
        "advanced_review": {"enabled": True},
        "scope_review": {"enabled": False, "disallowed_terms": []},
        "content_scope": {"enabled": False, "forbidden_terms": []},
    }


def _finding_types(result: dict) -> set[str]:
    return {str(finding.get("type")) for finding in result["raw_findings"]}


def _visible_text(doc: Document) -> dict:
    return {
        "body_paragraphs": [paragraph.text for paragraph in doc.paragraphs],
        "tables": [
            [
                [
                    [paragraph.text for paragraph in cell.paragraphs]
                    for cell in row.cells
                ]
                for row in table.rows
            ]
            for table in doc.tables
        ],
        "headers": [
            [paragraph.text for paragraph in section.header.paragraphs]
            for section in doc.sections
        ],
        "footers": [
            [paragraph.text for paragraph in section.footer.paragraphs]
            for section in doc.sections
        ],
    }


def _collect_user_facing_report_text(value) -> str:
    parts: list[str] = []
    if isinstance(value, dict):
        for key, child in value.items():
            if key in {"issue_id", "rule_id", "group_id"}:
                continue
            parts.append(_collect_user_facing_report_text(child))
    elif isinstance(value, list):
        for child in value:
            parts.append(_collect_user_facing_report_text(child))
    elif isinstance(value, str):
        parts.append(value)
    return "\n".join(part for part in parts if part)


def _assert_no_banned_user_facing_text(text: str) -> None:
    banned_fragments = [
        "Set alignment",
        "Set font",
        "Set line spacing",
        "does not match the required format",
        "issue(s)",
        "occurrence(s)",
        "PAGE_",
        "margin_top_cm",
        "margin_bottom_cm",
        "Dong trong lien tiep",
        "Co nhieu dong",
        "Kiem tra cac dong",
        "pages, excluding appendix",
        "on one centered line",
        "Khong vuot qua",
    ]
    for fragment in banned_fragments:
        if fragment in text:
            raise AssertionError(f"Unexpected raw user-facing text fragment: {fragment}")


if __name__ == "__main__":
    unittest.main()
