from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from app.services.docx_formatter.engine.annotation_engine import AnnotationEngine
from app.services.docx_formatter.factories.rule_factory import RuleFactory
from app.services.docx_formatter.rules.caption_format_rule import CaptionFormatRule
from app.services.docx_formatter.rules.caption_numbering_rule import CaptionNumberingRule
from app.services.docx_formatter.rules.character_density_rule import CharacterDensityRule
from app.services.docx_formatter.rules.front_matter_heading_rule import FrontMatterHeadingRule
from app.services.docx_formatter.rules.header_footer_format_rule import HeaderFooterFormatRule
from app.services.docx_formatter.rules.heading_format_rule import HeadingFormatRule
from app.services.docx_formatter.rules.list_item_format_rule import ListItemFormatRule
from app.services.docx_formatter.rules.page_numbering_rule import PageNumberingRule
from app.services.docx_formatter.rules.page_setup_rule import PageSetupRule
from app.services.docx_formatter.rules.paragraph_format_rule import ParagraphFormatRule
from app.services.docx_formatter.rules.table_cell_format_rule import TableCellFormatRule
from app.services.docx_formatter.utils.submission_cleanup import clean_submission_artifacts
from app.services.docx_formatter.writers.word_comment_writer import WordCommentWriter


class PipelineSplitTests(unittest.TestCase):
    def test_analyze_and_fix_registries_are_separate(self) -> None:
        analyze_rule_types = {type(rule) for rule in RuleFactory.create_analyze_rules()}
        fix_rule_types = {type(rule) for rule in RuleFactory.create_fix_rules()}

        self.assertIn(CaptionNumberingRule, analyze_rule_types)
        self.assertIn(CharacterDensityRule, analyze_rule_types)
        self.assertEqual(
            fix_rule_types,
            {
                PageSetupRule,
                FrontMatterHeadingRule,
                ListItemFormatRule,
                CaptionFormatRule,
                TableCellFormatRule,
                HeaderFooterFormatRule,
                ParagraphFormatRule,
                HeadingFormatRule,
            },
        )
        self.assertNotIn(CaptionNumberingRule, fix_rule_types)
        self.assertNotIn(CharacterDensityRule, fix_rule_types)
        self.assertNotIn(PageNumberingRule, fix_rule_types)

    def test_annotation_groups_multiple_findings_on_same_paragraph(self) -> None:
        findings = [
            _finding("PARAGRAPH_ALIGNMENT_ERROR", 2, "alignment"),
            _finding("PARAGRAPH_FONT_SIZE_ERROR", 2, "font_size"),
        ]

        comments = AnnotationEngine().build_comments(findings)

        self.assertEqual(len(comments), 1)
        self.assertEqual(comments[0].source_count, 2)

    def test_word_comment_writer_reports_grouped_comment_stats(self) -> None:
        if not hasattr(Document(), "add_comment"):
            self.skipTest("python-docx runtime does not support Document.add_comment")

        findings = [
            _finding("PARAGRAPH_ALIGNMENT_ERROR", 1, "alignment"),
            _finding("PARAGRAPH_FONT_SIZE_ERROR", 1, "font_size"),
        ]
        comments = AnnotationEngine().build_comments(findings)

        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "annotated.docx"
            doc = Document()
            doc.add_paragraph("Đoạn văn dùng để gắn comment.")
            doc.save(input_path)

            result = WordCommentWriter().write(
                input_path=str(input_path),
                output_path=str(output_path),
                comments=comments,
            )

            self.assertEqual(result["total_findings"], 2)
            self.assertEqual(result["total_comments_created"], 1)
            self.assertEqual(result["skipped_findings"], 0)
            self.assertIn("gom", result["comment_note"])

    def test_submission_cleanup_removes_comments_and_highlight(self) -> None:
        if not hasattr(Document(), "add_comment"):
            self.skipTest("python-docx runtime does not support Document.add_comment")

        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "fixed.docx"
            doc = Document()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("Đoạn có đánh dấu nháp cần dọn.")
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            doc.add_comment([run], text="Comment nháp", author="Test", initials="TT")
            doc.save(docx_path)

            clean_submission_artifacts(str(docx_path))

            with zipfile.ZipFile(docx_path, "r") as docx_zip:
                names = set(docx_zip.namelist())
                document_xml = docx_zip.read("word/document.xml")

            self.assertNotIn("word/comments.xml", names)
            self.assertNotIn(b"commentRangeStart", document_xml)
            self.assertNotIn(b"commentReference", document_xml)
            self.assertNotIn(b"<w:highlight", document_xml)
            self.assertNotIn(b"CommentReference", document_xml)


def _finding(finding_type: str, paragraph_index: int, field: str) -> dict:
    return {
        "type": finding_type,
        "severity": "warning",
        "location": f"Paragraph {paragraph_index}",
        "message": "Formatting issue.",
        "current_value": "Sai",
        "expected_value": "Đúng",
        "suggestion": "Kiểm tra định dạng.",
        "metadata": {
            "target": "paragraph",
            "context": "body_paragraph",
            "report_group_id": "body_paragraph",
            "paragraph_index": paragraph_index,
            "field": field,
            "auto_fixable": True,
            "manual_review": False,
            "fix_action": {"type": "set_formatting"},
        },
    }


if __name__ == "__main__":
    unittest.main()
