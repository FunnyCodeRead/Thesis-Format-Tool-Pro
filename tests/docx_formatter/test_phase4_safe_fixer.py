from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt

from app.services.docx_formatter.analyzer import analyze_document_with_details
from app.services.docx_formatter.fixer import fix_document


class Phase4SafeFixerTests(unittest.TestCase):
    def test_fix_document_preserves_text_and_cleans_submission_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "fixed.docx"

            doc = Document()
            body = doc.add_paragraph()
            body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = body.add_run(
                "Doan van noi dung du dai de duoc xem la body paragraph va can duoc can deu hai ben."
            )
            run.font.name = "Arial"
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            if hasattr(doc, "add_comment"):
                doc.add_comment([run], text="Comment nhap phai duoc xoa khoi file nop.", author="Test")

            list_paragraph = doc.add_paragraph("- Muc bullet/list khong bi ep first-line indent nhu body paragraph.")
            list_paragraph.paragraph_format.first_line_indent = None

            caption = doc.add_paragraph("Hinh 1.1. So do tong quan")
            caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).paragraphs[0].add_run("Noi dung bang khong bi sua text.")

            footer_run = doc.sections[0].footer.paragraphs[0].add_run("Footer text giu nguyen.")
            footer_run.font.name = "Arial"
            doc.save(input_path)

            original_text = _visible_text(Document(input_path))

            result = fix_document(
                input_path=str(input_path),
                document_type="do_an_tot_nghiep",
                output_path=str(output_path),
            )

            fixed_doc = Document(output_path)
            self.assertEqual(_visible_text(fixed_doc), original_text)
            self.assertEqual(result["fix_mode"], "safe_all")
            self.assertEqual(
                result["safe_fix_rules"],
                [
                    "PageSetupRule",
                    "FrontMatterHeadingRule",
                    "ListItemFormatRule",
                    "CaptionFormatRule",
                    "TableCellFormatRule",
                    "HeaderFooterFormatRule",
                    "ParagraphFormatRule",
                    "HeadingFormatRule",
                ],
            )
            self.assertIn("CaptionNumberingRule", result["blocked_fix_rules"])
            self.assertIn("PageNumberingRule", result["blocked_fix_rules"])
            self.assertIn("Font chữ và cỡ chữ", result["safe_fix_scope"])
            self.assertEqual(result["changes_by_rule"]["ParagraphFormatRule"], result["paragraph_changes"])
            self.assertEqual(result["changes_by_rule"]["ListItemFormatRule"], result["list_item_changes"])
            self.assertEqual(result["changes_by_rule"]["CaptionFormatRule"], result["caption_changes"])
            self.assertEqual(result["changes_by_rule"]["HeaderFooterFormatRule"], result["header_footer_changes"])
            self.assertTrue(result["safety_checks"]["original_not_overwritten"])
            self.assertTrue(result["safety_checks"]["visible_text_preserved"])
            self.assertTrue(result["safety_checks"]["comments_removed"])
            self.assertTrue(result["safety_checks"]["highlights_removed"])
            self.assertTrue(result["safety_checks"]["app_error_markers_not_added"])
            self.assertEqual(result["cleanup_report"]["after"]["comment_parts"], 0)
            self.assertEqual(result["cleanup_report"]["after"]["comment_markers"], 0)
            self.assertEqual(result["cleanup_report"]["after"]["highlights"], 0)
            self.assertIsNone(fixed_doc.paragraphs[1].paragraph_format.first_line_indent)
            self.assertEqual(fixed_doc.paragraphs[2].paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(fixed_doc.sections[0].footer.paragraphs[0].runs[0].font.name, "Times New Roman")

            with zipfile.ZipFile(output_path, "r") as docx_zip:
                names = set(docx_zip.namelist())
                document_xml = docx_zip.read("word/document.xml")

            self.assertNotIn("word/comments.xml", names)
            self.assertNotIn(b"commentRangeStart", document_xml)
            self.assertNotIn(b"commentReference", document_xml)
            self.assertNotIn(b"<w:highlight", document_xml)

    def test_fix_document_rejects_overwriting_original(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            doc = Document()
            doc.add_paragraph("Doan van test.")
            doc.save(input_path)

            with self.assertRaises(Exception):
                fix_document(
                    input_path=str(input_path),
                    document_type="do_an_tot_nghiep",
                    output_path=str(input_path),
                )

    def test_fix_document_formats_list_items_and_captions_safely(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "fixed.docx"

            doc = Document()
            list_paragraph = doc.add_paragraph("- Muc danh sach can duoc dinh dang spacing an toan.")
            list_paragraph.paragraph_format.space_before = Pt(5)
            list_paragraph.paragraph_format.space_after = Pt(12)
            list_paragraph.paragraph_format.line_spacing = 1.0

            caption = doc.add_paragraph("Hinh 1.1. So do tong quan")
            caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            caption.paragraph_format.space_before = Pt(0)
            caption.paragraph_format.space_after = Pt(12)
            doc.save(input_path)

            result = fix_document(
                input_path=str(input_path),
                document_type="do_an_tot_nghiep",
                output_path=str(output_path),
            )
            analysis = analyze_document_with_details(str(output_path), "do_an_tot_nghiep")
            remaining_types = {finding["type"] for finding in analysis["raw_findings"]}

            self.assertGreater(result["list_item_changes"], 0)
            self.assertGreater(result["caption_changes"], 0)
            self.assertNotIn("LIST_ITEM_SPACE_BEFORE_PT_ERROR", remaining_types)
            self.assertNotIn("LIST_ITEM_SPACE_AFTER_PT_ERROR", remaining_types)
            self.assertNotIn("LIST_ITEM_LINE_SPACING_ERROR", remaining_types)
            self.assertNotIn("CAPTION_SPACE_BEFORE_PT_ERROR", remaining_types)
            self.assertNotIn("CAPTION_SPACE_AFTER_PT_ERROR", remaining_types)
            self.assertNotIn("CAPTION_ALIGNMENT_ERROR", remaining_types)


def _visible_text(doc: Document) -> dict:
    return {
        "body_paragraphs": [paragraph.text for paragraph in doc.paragraphs],
        "tables": [
            [
                [
                    [paragraph.text for paragraph in cell.paragraphs]
                    for cell in row.cells
                ]
                for row in table.rows
            ]
            for table in doc.tables
        ],
        "headers": [
            [paragraph.text for paragraph in section.header.paragraphs]
            for section in doc.sections
        ],
        "footers": [
            [paragraph.text for paragraph in section.footer.paragraphs]
            for section in doc.sections
        ],
    }


if __name__ == "__main__":
    unittest.main()
