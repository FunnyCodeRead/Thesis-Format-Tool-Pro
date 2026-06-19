from __future__ import annotations

import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services.docx_formatter.engine.caption_detector import detect_caption_text
from app.services.docx_formatter.engine.context_classifier import classify_paragraph_context
from app.services.docx_formatter.rules.caption_numbering_rule import CaptionNumberingRule
from app.services.docx_formatter.rules.heading_structure_rule import HeadingStructureRule
from app.services.docx_formatter.rules.page_numbering_rule import PageNumberingRule
from app.services.docx_formatter.rules.paragraph_format_rule import ParagraphFormatRule
from app.services.docx_formatter.rules.scope_review_rule import ScopeReviewRule


class AnalyzerHardeningTests(unittest.TestCase):
    def test_caption_detector_separates_caption_body_reference_and_malformed(self) -> None:
        self.assertEqual(detect_caption_text("Bảng 1.1 trình bày dữ liệu thử nghiệm.").status, "body_reference")
        self.assertEqual(detect_caption_text("Bảng 1.1. Thông số thử nghiệm").status, "valid")
        self.assertEqual(detect_caption_text("Hình 2.2.3.3. Số bị rối").status, "malformed")
        self.assertEqual(detect_caption_text("Hình 2.1 Mô hình tổng quan").status, "missing_separator")

    def test_body_reference_is_not_classified_as_caption(self) -> None:
        doc = Document()
        paragraph = doc.add_paragraph("Bảng 1.1 trình bày dữ liệu thử nghiệm trong chương.")

        context = classify_paragraph_context(paragraph, 1)

        self.assertNotEqual(context.context, "caption")

    def test_caption_numbering_reports_malformed_and_list_duplicate(self) -> None:
        doc = Document()
        doc.add_paragraph("Chương 1.")
        doc.add_paragraph("Hình 2.2.3.3. Số hình bị rối")
        doc.add_paragraph("DANH MỤC BẢNG")
        doc.add_paragraph("Bảng 3.28. Thống kê lỗi ........ 10")
        doc.add_paragraph("Bảng 3.28. Thống kê lỗi ........ 11")

        findings = CaptionNumberingRule().analyze(doc, _base_config())
        finding_types = {finding.type for finding in findings}

        self.assertIn("FIGURE_NUMBERING_MALFORMED_REVIEW", finding_types)
        self.assertIn("TABLE_LIST_DUPLICATE_REVIEW", finding_types)

    def test_heading_number_mismatch_uses_nearest_chapter(self) -> None:
        doc = Document()
        doc.add_paragraph("Chương 1.")
        doc.add_paragraph("2.2 Cơ sở lý thuyết")

        findings = HeadingStructureRule().analyze(doc, _base_config())

        self.assertEqual([finding.type for finding in findings], ["HEADING_CHAPTER_MISMATCH_REVIEW"])
        self.assertIn("Chương 1", findings[0].current_value or "")

    def test_bullet_list_is_not_checked_as_body_paragraph(self) -> None:
        doc = Document()
        doc.add_paragraph("- Một mục gạch đầu dòng trong nội dung báo cáo")

        findings = ParagraphFormatRule().analyze(doc, _base_config())

        self.assertFalse(any(finding.type.startswith("PARAGRAPH_") for finding in findings))

    def test_page_number_unknown_alignment_message_is_vietnamese(self) -> None:
        doc = Document()
        footer_paragraph = doc.sections[0].footer.paragraphs[0]
        _add_page_field(footer_paragraph)

        findings = PageNumberingRule().analyze(doc, _base_config())
        alignment_findings = [
            finding for finding in findings if finding.type == "PAGE_NUMBER_ALIGNMENT_REVIEW"
        ]

        self.assertEqual(len(alignment_findings), 1)
        self.assertIn("không có căn lề trực tiếp", alignment_findings[0].current_value or "")
        self.assertEqual(alignment_findings[0].metadata.get("part_paragraph_index"), 1)

    def test_content_scope_is_disabled_by_default_and_enabled_by_override(self) -> None:
        doc = Document()
        doc.add_paragraph("Safe Route là cụm từ chỉ nên được kiểm tra khi project bật scope riêng.")

        self.assertEqual(ScopeReviewRule().analyze(doc, _base_config()), [])

        config = _base_config()
        config["content_scope"] = {
            "enabled": True,
            "forbidden_terms": ["Safe Route"],
        }
        findings = ScopeReviewRule().analyze(doc, config)

        self.assertEqual([finding.type for finding in findings], ["OUT_OF_SCOPE_TERM_REVIEW"])


def _base_config() -> dict:
    return {
        "paragraph": {
            "font_name": "Times New Roman",
            "font_size": 13,
            "line_spacing": 1.3,
            "alignment": "JUSTIFY",
            "first_line_indent_cm": 1.0,
            "space_before_pt": 6,
            "space_after_pt": 6,
        },
        "chapter_layout": {"enabled": True},
        "caption_numbering": {
            "enabled": True,
            "min_number_parts": 2,
            "max_number_parts": 2,
            "require_chapter_prefix": True,
        },
        "pagination": {"enabled": True},
        "scope_review": {"enabled": False, "disallowed_terms": []},
        "content_scope": {"enabled": False, "forbidden_terms": []},
    }


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


if __name__ == "__main__":
    unittest.main()
