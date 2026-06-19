from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.services.docx_formatter.fixer import fix_document


class StyleLevelFixerTests(unittest.TestCase):
    def test_fix_document_updates_exclusive_body_style_before_paragraph_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "fixed.docx"

            doc = Document()
            style = doc.styles.add_style("Body Safe Test", WD_STYLE_TYPE.PARAGRAPH)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style.font.name = "Arial"
            style.font.size = Pt(12)

            for index in range(3):
                paragraph = doc.add_paragraph(
                    f"Day la doan van noi dung thu {index + 1} du dai de duoc xem la body paragraph va can dung style rieng.",
                    style=style,
                )
                paragraph.add_run(" Noi dung giu nguyen sau khi sua style.")

            doc.save(input_path)

            result = fix_document(
                input_path=str(input_path),
                document_type="do_an_tot_nghiep",
                output_path=str(output_path),
            )
            fixed_doc = Document(output_path)
            fixed_style = fixed_doc.styles["Body Safe Test"]

            self.assertGreater(result["style_changes"], 0)
            self.assertEqual(result["style_fix_groups_applied"], 1)
            self.assertEqual(result["changes_by_rule"]["StyleLevelFixEngine"], result["style_changes"])
            self.assertEqual(fixed_style.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assertEqual(fixed_style.font.name, "Times New Roman")
            self.assertEqual(round(fixed_style.font.size.pt, 1), 13.0)
            self.assertTrue(result["safety_checks"]["visible_text_preserved"])

    def test_fix_document_skips_style_level_when_style_is_shared_with_cover(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.docx"
            output_path = Path(temp_dir) / "fixed.docx"

            doc = Document()
            style = doc.styles.add_style("Shared Cover Body Test", WD_STYLE_TYPE.PARAGRAPH)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style.font.name = "Arial"
            style.font.size = Pt(12)

            doc.add_paragraph("TRUONG DAI HOC CONG NGHE THONG TIN", style=style)
            for index in range(3):
                doc.add_paragraph(
                    f"Day la doan van noi dung thu {index + 1} du dai de test style dung lan giua bia va body.",
                    style=style,
                )
            doc.save(input_path)

            result = fix_document(
                input_path=str(input_path),
                document_type="do_an_tot_nghiep",
                output_path=str(output_path),
            )
            fixed_doc = Document(output_path)
            fixed_style = fixed_doc.styles["Shared Cover Body Test"]
            skipped_style_names = {
                str(item.get("style_name"))
                for item in result["style_fix_skipped"]
                if item.get("reason")
            }

            self.assertEqual(result["style_changes"], 0)
            self.assertIn("Shared Cover Body Test", skipped_style_names)
            self.assertEqual(fixed_style.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertEqual(fixed_style.font.name, "Arial")
            self.assertEqual(round(fixed_style.font.size.pt, 1), 12.0)
            self.assertTrue(result["paragraph_changes"] > 0)
            self.assertTrue(result["safety_checks"]["visible_text_preserved"])


if __name__ == "__main__":
    unittest.main()
