from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from app.services.docx_formatter.analyzer import (
    DocumentAnalysisError,
    analyze_document_with_details,
)
from app.services.docx_formatter.config.config_loader import AnalyzerConfigError, ConfigLoader
from app.services.docx_formatter.engine.fixer_engine import FixerEngine
from app.services.docx_formatter.engine.style_level_fixer import StyleLevelFixEngine
from app.services.docx_formatter.factories.rule_factory import RuleFactory
from app.services.docx_formatter.utils.submission_cleanup import (
    clean_submission_artifacts,
    has_submission_artifacts,
    has_tracked_changes,
    inspect_submission_artifacts,
    inspect_tracked_changes,
)

FIX_MODE_SAFE_ALL = "safe_all"
FIX_MODE_SAFE_SCOPE = "safe_scope"

SAFE_FIX_SCOPE_REGISTRY: dict[str, dict[str, Any]] = {
    "page_setup": {
        "label": "Căn lề trang và khổ giấy A4",
        "rule": "PageSetupRule",
        "style_contexts": [],
    },
    "front_matter_heading": {
        "label": "Định dạng tiêu đề phần đầu an toàn",
        "rule": "FrontMatterHeadingRule",
        "style_contexts": ["front_matter_heading"],
    },
    "list_item": {
        "label": "Định dạng bullet/list an toàn",
        "rule": "ListItemFormatRule",
        "style_contexts": ["list_item"],
    },
    "caption_format": {
        "label": "Định dạng caption an toàn",
        "rule": "CaptionFormatRule",
        "style_contexts": ["caption"],
    },
    "table_cell_format": {
        "label": "Định dạng font/cỡ chữ trong ô bảng",
        "rule": "TableCellFormatRule",
        "style_contexts": [],
    },
    "header_footer_format": {
        "label": "Định dạng font/cỡ chữ header/footer",
        "rule": "HeaderFooterFormatRule",
        "style_contexts": [],
    },
    "paragraph_format": {
        "label": "Font, cỡ chữ, căn đoạn, thụt đầu dòng, giãn dòng và khoảng cách đoạn",
        "rule": "ParagraphFormatRule",
        "style_contexts": ["body_paragraph"],
    },
    "heading_format": {
        "label": "Định dạng heading cơ bản",
        "rule": "HeadingFormatRule",
        "style_contexts": ["heading"],
    },
}

SAFE_FIX_SCOPE = [
    "Font chữ và cỡ chữ",
    *[item["label"] for item in SAFE_FIX_SCOPE_REGISTRY.values()],
]


class DocumentFixError(RuntimeError):
    pass


def fix_document(
    input_path: str,
    document_type: str,
    output_path: str,
    *,
    config_override: dict[str, Any] | None = None,
    fix_options: dict[str, Any] | None = None,
) -> dict[str, Any]:
    if Path(input_path).resolve() == Path(output_path).resolve():
        raise DocumentFixError("Output path must be different from the original input path.")

    tracked_changes_report = inspect_tracked_changes(input_path)
    if has_tracked_changes(tracked_changes_report):
        raise DocumentFixError(
            "File còn Track Changes thật. Hãy accept/reject toàn bộ thay đổi trong Word, "
            "lưu lại file sạch rồi upload lại trước khi dùng chức năng tự sửa định dạng."
        )

    try:
        config = ConfigLoader().load(document_type, config_override=config_override)
        doc = Document(input_path)
    except AnalyzerConfigError:
        raise
    except Exception as exc:
        raise DocumentFixError("Failed to load .docx document.") from exc

    fix_plan = _normalize_fix_options(fix_options)
    original_text_snapshot = _visible_text_snapshot(doc)
    all_fix_rules = RuleFactory.create_fix_rules()
    fix_rules = _select_fix_rules(all_fix_rules, fix_plan["scope_keys"])
    analyze_rules = RuleFactory.create_analyze_rules()
    safe_fix_rules = [rule.__class__.__name__ for rule in fix_rules]
    all_safe_fix_rules = [rule.__class__.__name__ for rule in all_fix_rules]
    blocked_fix_rules = [
        rule.__class__.__name__
        for rule in analyze_rules
        if rule.__class__.__name__ not in set(all_safe_fix_rules)
    ]
    style_fix_result = StyleLevelFixEngine().fix(
        doc,
        config,
        allowed_contexts=_selected_style_contexts(fix_plan["scope_keys"]),
    )
    engine = FixerEngine(fix_rules)
    result = engine.fix(doc, config)
    style_changes = int(style_fix_result["style_changes"])
    result["style_changes"] = style_changes
    result["style_fix_mode"] = style_fix_result["style_fix_mode"]
    result["style_fix_groups_applied"] = style_fix_result["style_fix_groups_applied"]
    result["style_changes_by_style"] = style_fix_result["style_changes_by_style"]
    result["style_fix_skipped"] = style_fix_result["style_fix_skipped"]
    result["total_changes"] += style_changes
    result["changes_by_rule"]["StyleLevelFixEngine"] = style_changes

    try:
        doc.save(output_path)
        artifacts_before_cleanup = inspect_submission_artifacts(output_path)
        clean_submission_artifacts(output_path)
        artifacts_after_cleanup = inspect_submission_artifacts(output_path)
        if has_submission_artifacts(artifacts_after_cleanup):
            raise DocumentFixError("Fixed .docx still contains comment or highlight artifacts.")
        fixed_doc = Document(output_path)
    except Exception as exc:
        if isinstance(exc, DocumentFixError):
            raise
        raise DocumentFixError("Failed to save fixed .docx document.") from exc

    fixed_text_snapshot = _visible_text_snapshot(fixed_doc)
    visible_text_preserved = fixed_text_snapshot == original_text_snapshot
    if not visible_text_preserved:
        raise DocumentFixError(
            "Fixed .docx failed safety validation because visible document text changed."
        )

    post_fix_validation = _validate_fixed_output(
        output_path=output_path,
        document_type=document_type,
        config_override=config_override,
        selected_scope=fix_plan["scope_keys"],
    )
    if post_fix_validation["status"] != "passed":
        sample_types = ", ".join(post_fix_validation["blocking_safe_issue_types"][:5])
        raise DocumentFixError(
            "Post-fix analyzer gate failed because selected safe formatting issues remain"
            + (f": {sample_types}." if sample_types else ".")
        )

    return {
        **result,
        "output_path": output_path,
        "fix_mode": fix_plan["fix_mode"],
        "applied_fix_scope": fix_plan["scope_keys"],
        "available_fix_scope": _available_fix_scope(),
        "safe_fix_rules": safe_fix_rules,
        "skipped_safe_fix_rules": [
            rule_name for rule_name in all_safe_fix_rules if rule_name not in set(safe_fix_rules)
        ],
        "blocked_fix_rules": blocked_fix_rules,
        "safe_fix_scope": SAFE_FIX_SCOPE,
        "safety_checks": {
            "original_not_overwritten": Path(input_path).resolve() != Path(output_path).resolve(),
            "visible_text_preserved": visible_text_preserved,
            "source_is_original_file": True,
            "comments_removed": artifacts_after_cleanup["comment_parts"] == 0
            and artifacts_after_cleanup["comment_relationships"] == 0
            and artifacts_after_cleanup["comment_markers"] == 0
            and artifacts_after_cleanup["comment_reference_runs"] == 0,
            "highlights_removed": artifacts_after_cleanup["highlights"] == 0,
            "app_error_markers_not_added": _app_error_marker_count(fixed_text_snapshot)
            == _app_error_marker_count(original_text_snapshot),
        },
        "cleanup_report": {
            "before": artifacts_before_cleanup,
            "after": artifacts_after_cleanup,
        },
        "post_fix_validation": post_fix_validation,
    }


def _normalize_fix_options(fix_options: dict[str, Any] | None) -> dict[str, Any]:
    options = fix_options or {}
    fix_mode = str(options.get("fix_mode") or FIX_MODE_SAFE_ALL).strip()
    if fix_mode not in {FIX_MODE_SAFE_ALL, FIX_MODE_SAFE_SCOPE}:
        raise DocumentFixError(
            "fix_mode không hợp lệ. Chỉ hỗ trợ safe_all hoặc safe_scope."
        )

    if fix_mode == FIX_MODE_SAFE_ALL:
        scope_keys = list(SAFE_FIX_SCOPE_REGISTRY.keys())
    else:
        raw_scope = options.get("fix_scope") or []
        if not isinstance(raw_scope, list):
            raise DocumentFixError("fix_scope phải là danh sách nhóm sửa an toàn.")

        scope_keys = []
        for item in raw_scope:
            key = str(item).strip()
            if not key:
                continue
            scope_keys.append(key)

        unknown_scope = [
            key for key in scope_keys if key not in SAFE_FIX_SCOPE_REGISTRY
        ]
        if unknown_scope:
            allowed = ", ".join(SAFE_FIX_SCOPE_REGISTRY.keys())
            raise DocumentFixError(
                f"fix_scope không hợp lệ: {', '.join(unknown_scope)}. "
                f"Các nhóm được hỗ trợ: {allowed}."
            )

        scope_keys = list(dict.fromkeys(scope_keys))
        if not scope_keys:
            raise DocumentFixError("fix_scope cần ít nhất một nhóm sửa an toàn.")

    return {
        "fix_mode": fix_mode,
        "scope_keys": scope_keys,
    }


def _select_fix_rules(rules: list[Any], scope_keys: list[str]) -> list[Any]:
    allowed_rule_names = {
        str(SAFE_FIX_SCOPE_REGISTRY[key]["rule"])
        for key in scope_keys
        if SAFE_FIX_SCOPE_REGISTRY[key].get("rule")
    }
    return [rule for rule in rules if rule.__class__.__name__ in allowed_rule_names]


def _selected_style_contexts(scope_keys: list[str]) -> set[str]:
    contexts: set[str] = set()
    for key in scope_keys:
        contexts.update(str(item) for item in SAFE_FIX_SCOPE_REGISTRY[key]["style_contexts"])
    return contexts


def _available_fix_scope() -> list[dict[str, str]]:
    return [
        {
            "key": key,
            "label": str(value["label"]),
            "rule": str(value["rule"]),
        }
        for key, value in SAFE_FIX_SCOPE_REGISTRY.items()
    ]


def _validate_fixed_output(
    *,
    output_path: str,
    document_type: str,
    config_override: dict[str, Any] | None,
    selected_scope: list[str],
) -> dict[str, Any]:
    try:
        analysis = analyze_document_with_details(
            output_path,
            document_type,
            config_override=config_override,
        )
    except (AnalyzerConfigError, DocumentAnalysisError) as exc:
        raise DocumentFixError("Post-fix analyzer gate could not analyze the fixed file.") from exc

    raw_findings = analysis.get("raw_findings", [])
    selected_scope_set = set(selected_scope)
    safe_findings = [
        finding
        for finding in raw_findings
        if _fixability_scope(finding) == "safe_auto_fix"
    ]
    blocking_findings = [
        finding
        for finding in safe_findings
        if _finding_scope_key(finding) in selected_scope_set
    ]
    unselected_safe_findings = [
        finding
        for finding in safe_findings
        if _finding_scope_key(finding) not in selected_scope_set
    ]
    manual_review_findings = [
        finding
        for finding in raw_findings
        if _fixability_scope(finding) == "manual_review"
    ]

    return {
        "status": "passed" if not blocking_findings else "failed",
        "checked": True,
        "selected_scope": selected_scope,
        "remaining_total_findings": len(raw_findings),
        "remaining_safe_auto_fix_count": len(safe_findings),
        "remaining_selected_safe_issue_count": len(blocking_findings),
        "remaining_unselected_safe_issue_count": len(unselected_safe_findings),
        "remaining_manual_review_count": len(manual_review_findings),
        "blocking_safe_issue_count": len(blocking_findings),
        "blocking_safe_issue_types": _finding_type_counts(blocking_findings),
        "blocking_safe_issue_samples": _finding_samples(blocking_findings),
        "render_verification": analysis.get("render_verification", {}),
    }


def _fixability_scope(finding: dict[str, Any]) -> str:
    metadata = finding.get("metadata")
    if isinstance(metadata, dict):
        return str(metadata.get("fixability_scope") or "")
    return ""


def _finding_scope_key(finding: dict[str, Any]) -> str:
    metadata = finding.get("metadata")
    metadata = metadata if isinstance(metadata, dict) else {}
    group_id = str(metadata.get("report_group_id") or "")
    finding_type = str(finding.get("type") or "").upper()

    if finding_type in {"PAGE_MARGIN_ERROR", "PAPER_SIZE_ERROR"} or group_id == "page_setup":
        return "page_setup"
    if finding_type.startswith("FRONT_MATTER_HEADING_") or group_id == "front_matter":
        return "front_matter_heading"
    if finding_type.startswith("LIST_ITEM_") or group_id == "list_item":
        return "list_item"
    if finding_type.startswith("CAPTION_") or group_id == "caption":
        return "caption_format"
    if finding_type.startswith("TABLE_CELL_") or group_id == "table_cell_format":
        return "table_cell_format"
    if finding_type.startswith("HEADER_FOOTER_") or group_id == "header_footer_format":
        return "header_footer_format"
    if finding_type.startswith("PARAGRAPH_") or group_id == "body_paragraph":
        return "paragraph_format"
    if finding_type.startswith("HEADING_") or group_id == "heading":
        return "heading_format"
    return group_id or "unknown"


def _finding_type_counts(findings: list[dict[str, Any]]) -> list[str]:
    counts: dict[str, int] = {}
    for finding in findings:
        finding_type = str(finding.get("type") or "UNKNOWN")
        counts[finding_type] = counts.get(finding_type, 0) + 1
    return [
        f"{finding_type} ({count})"
        for finding_type, count in sorted(counts.items(), key=lambda item: (-item[1], item[0]))
    ]


def _finding_samples(findings: list[dict[str, Any]]) -> list[dict[str, Any]]:
    samples: list[dict[str, Any]] = []
    for finding in findings[:10]:
        samples.append(
            {
                "type": finding.get("type"),
                "location": finding.get("location"),
                "message": finding.get("message"),
                "scope": _finding_scope_key(finding),
            }
        )
    return samples


def _visible_text_snapshot(doc: Any) -> dict[str, Any]:
    return {
        "body_paragraphs": [paragraph.text for paragraph in doc.paragraphs],
        "tables": [
            [
                [
                    [paragraph.text for paragraph in cell.paragraphs]
                    for cell in row.cells
                ]
                for row in table.rows
            ]
            for table in doc.tables
        ],
        "headers": [
            [paragraph.text for paragraph in section.header.paragraphs]
            for section in doc.sections
        ],
        "footers": [
            [paragraph.text for paragraph in section.footer.paragraphs]
            for section in doc.sections
        ],
    }


def _app_error_marker_count(snapshot: dict[str, Any]) -> int:
    serialized = repr(snapshot)
    return serialized.count("[LỖI]") + serialized.count("[LOI]")
