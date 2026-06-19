from __future__ import annotations

import re
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.services.docx_formatter.domain.finding import Finding
from app.services.docx_formatter.domain.rule import AnalyzeRule
from app.services.docx_formatter.engine.context_classifier import text_preview
from app.services.docx_formatter.rules.common_format import CommonFormatMixin
from app.services.docx_formatter.utils.text_utils import normalize_text

ROMAN_FORMATS = {"lowerRoman", "upperRoman"}
DECIMAL_FORMATS = {"decimal", None}


class PageNumberingRule(CommonFormatMixin, AnalyzeRule):
    def analyze(self, doc: Any, config: dict[str, Any]) -> list[Finding]:
        pagination_config = config.get("pagination", {})
        if pagination_config.get("enabled", True) is False:
            return []

        findings: list[Finding] = []
        findings.extend(self._cover_numbering_findings(doc))
        findings.extend(self._footer_alignment_findings(doc))
        findings.extend(self._roman_restart_findings(doc))
        findings.extend(self._main_content_reset_findings(doc))
        return findings

    def _cover_numbering_findings(self, doc: Any) -> list[Finding]:
        if not doc.sections:
            return []

        first_section = doc.sections[0]
        if not _part_has_page_field(first_section.footer):
            return []

        if bool(getattr(first_section, "different_first_page_header_footer", False)):
            first_page_footer = getattr(first_section, "first_page_footer", None)
            if first_page_footer is None or not _part_has_page_field(first_page_footer):
                return []

        return [
            _manual_finding(
                finding_type="COVER_PAGE_NUMBER_VISIBLE_REVIEW",
                severity="warning",
                location="Section 1 Footer",
                message="Trang bìa có thể đang hiển thị số trang.",
                current_value=(
                    _part_text(first_section.footer)
                    or "Phát hiện PAGE field trong footer của section bìa; trang bìa có thể đang hiển thị số trang."
                ),
                expected_value="Trang bìa không hiển thị số trang.",
                suggestion="Bật Different First Page hoặc tách section để ẩn số trang ở trang bìa.",
                metadata={
                    "target": "footer",
                    "context": "header_footer",
                    "report_group_id": "header_footer_page_number",
                    "report_severity": "critical",
                    "section_index": 1,
                    "field": "page_number_position",
                    "text_preview": _part_text(first_section.footer),
                    "style_name": "Footer",
                },
            )
        ]

    def _footer_alignment_findings(self, doc: Any) -> list[Finding]:
        findings: list[Finding] = []
        for section_index, section in enumerate(doc.sections, start=1):
            for part_paragraph_index, paragraph in enumerate(section.footer.paragraphs, start=1):
                if not _paragraph_has_page_field(paragraph):
                    continue
                alignment = self._effective_alignment(paragraph)
                if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    continue
                if alignment is None:
                    message = f"Không xác định được vị trí căn số trang trong footer của Section {section_index}."
                    current_value = (
                        f"Section {section_index}, footer paragraph {part_paragraph_index}: "
                        "không có căn lề trực tiếp hoặc style kế thừa rõ ràng."
                    )
                    suggestion = (
                        "Mở footer của section này trong Word và kiểm tra paragraph chứa PAGE field; "
                        "nếu quy định yêu cầu thì đặt số trang căn giữa."
                    )
                else:
                    message = "Số trang trong footer chưa được căn giữa."
                    current_value = self._alignment_label(alignment)
                    suggestion = "Kiểm tra paragraph chứa PAGE field và căn giữa số trang ở footer theo quy định."
                findings.append(
                    _manual_finding(
                        finding_type="PAGE_NUMBER_ALIGNMENT_REVIEW",
                        severity="warning",
                        location=f"Section {section_index} Footer",
                        message=message,
                        current_value=current_value,
                        expected_value="Số trang căn giữa ở footer",
                        suggestion=suggestion,
                        metadata={
                            "target": "footer",
                            "context": "header_footer",
                            "report_group_id": "header_footer_page_number",
                            "report_severity": "major",
                            "section_index": section_index,
                            "part_paragraph_index": part_paragraph_index,
                            "field": "page_number_position",
                            "text_preview": text_preview(paragraph.text),
                            "style_name": "Footer",
                        },
                    )
                )
        return findings

    def _roman_restart_findings(self, doc: Any) -> list[Finding]:
        findings: list[Finding] = []
        roman_start_one_sections: list[int] = []

        for section_index, section in enumerate(doc.sections, start=1):
            settings = _page_number_settings(section)
            if settings["fmt"] in ROMAN_FORMATS and settings["start"] == "1":
                roman_start_one_sections.append(section_index)

        if len(roman_start_one_sections) <= 1:
            return findings

        for section_index in roman_start_one_sections[1:]:
            findings.append(
                _manual_finding(
                    finding_type="ROMAN_PAGE_NUMBER_REPEATED_REVIEW",
                    severity="warning",
                    location=f"Section {section_index}",
                    message="Số trang La Mã ở phần đầu có dấu hiệu bị restart/lặp lại.",
                    current_value=f"Section {section_index}: số La Mã đang bắt đầu lại từ i",
                    expected_value="Phần đầu chỉ restart số La Mã khi bắt đầu phần đánh số La Mã.",
                    suggestion="Kiểm tra lại section break và Page Number Format để tránh lặp số i.",
                    metadata={
                        "target": "section",
                        "context": "header_footer",
                        "report_group_id": "header_footer_page_number",
                        "report_severity": "major",
                        "section_index": section_index,
                        "field": "page_number_position",
                    },
                )
            )
        return findings

    def _main_content_reset_findings(self, doc: Any) -> list[Finding]:
        chapter_paragraph_index = _first_chapter_paragraph_index(doc)
        if chapter_paragraph_index is None:
            return []

        section_index = _section_index_for_paragraph(doc, chapter_paragraph_index)
        if section_index is None:
            return []

        section = doc.sections[section_index - 1]
        settings = _page_number_settings(section)
        findings: list[Finding] = []

        if settings["fmt"] not in DECIMAL_FORMATS:
            findings.append(
                _manual_finding(
                    finding_type="MAIN_PAGE_NUMBER_FORMAT_REVIEW",
                    severity="warning",
                    location=f"Section {section_index}",
                    message="Phần nội dung chính chưa dùng số trang Ả Rập.",
                    current_value=_page_number_format_label(settings["fmt"]),
                    expected_value="Số Ả Rập",
                    suggestion="Đặt phần nội dung chính về kiểu đánh số Ả Rập.",
                    metadata=_main_metadata(section_index, chapter_paragraph_index),
                )
            )

        if settings["start"] != "1":
            findings.append(
                _manual_finding(
                    finding_type="MAIN_PAGE_NUMBER_RESET_REVIEW",
                    severity="warning",
                    location=f"Section {section_index}",
                    message="Phần nội dung chính chưa reset số trang về 1.",
                    current_value=_page_number_start_label(settings["start"]),
                    expected_value="Bắt đầu từ 1",
                    suggestion="Kiểm tra Page Number Format của section nội dung chính và đặt Start at 1.",
                    metadata=_main_metadata(section_index, chapter_paragraph_index),
                )
            )

        return findings


def _main_metadata(section_index: int, paragraph_index: int) -> dict[str, Any]:
    return {
        "target": "section",
        "context": "header_footer",
        "report_group_id": "header_footer_page_number",
        "report_severity": "major",
        "section_index": section_index,
        "paragraph_index": paragraph_index,
        "field": "page_number_position",
    }


def _page_number_settings(section: Any) -> dict[str, str | None]:
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        return {"fmt": None, "start": None}
    return {
        "fmt": pg_num_type.get(qn("w:fmt")),
        "start": pg_num_type.get(qn("w:start")),
    }


def _page_number_format_label(value: str | None) -> str:
    if value is None:
        return "Chưa xác định kiểu đánh số trang trong section"
    labels = {
        "decimal": "Số Ả Rập",
        "lowerRoman": "Số La Mã thường",
        "upperRoman": "Số La Mã hoa",
    }
    return labels.get(value, value)


def _page_number_start_label(value: str | None) -> str:
    if value is None:
        return "Chưa thiết lập Start at cho section nội dung chính"
    return f"Bắt đầu từ {value}"


def _first_chapter_paragraph_index(doc: Any) -> int | None:
    for paragraph_index, paragraph in enumerate(doc.paragraphs, start=1):
        if re.match(r"^CHUONG\s+\d+\.?", normalize_text(paragraph.text.strip())):
            return paragraph_index
    return None


def _section_index_for_paragraph(doc: Any, paragraph_index: int) -> int | None:
    current_section = 1
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        if index == paragraph_index:
            return current_section
        if _paragraph_has_section_properties(paragraph):
            current_section += 1
    return min(current_section, len(doc.sections)) if doc.sections else None


def _paragraph_has_section_properties(paragraph: Any) -> bool:
    p_pr = getattr(paragraph._p, "pPr", None)
    return p_pr is not None and getattr(p_pr, "sectPr", None) is not None


def _part_has_page_field(part: Any) -> bool:
    return any(_paragraph_has_page_field(paragraph) for paragraph in part.paragraphs)


def _paragraph_has_page_field(paragraph: Any) -> bool:
    xml = getattr(getattr(paragraph, "_p", None), "xml", "")
    upper_xml = xml.upper()
    return "PAGE" in upper_xml or "NUMPAGES" in upper_xml


def _part_text(part: Any) -> str:
    text = " ".join(paragraph.text.strip() for paragraph in part.paragraphs if paragraph.text.strip())
    return text_preview(text)


def _manual_finding(
    *,
    finding_type: str,
    severity: str,
    location: str,
    message: str,
    current_value: str | None,
    expected_value: str | None,
    suggestion: str,
    metadata: dict[str, Any],
) -> Finding:
    return Finding(
        type=finding_type,
        severity=severity,
        location=location,
        message=message,
        current_value=current_value,
        expected_value=expected_value,
        suggestion=suggestion,
        metadata={
            **metadata,
            "auto_fixable": False,
            "manual_review": True,
            "fix_action": {
                "type": "manual_review",
                "reason": (
                    "Cần kiểm tra section/footer/page number thủ công; "
                    "hệ thống không tự sửa nội dung hoặc field phức tạp."
                ),
            },
        },
    )
