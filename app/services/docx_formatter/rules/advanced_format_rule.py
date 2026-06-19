from __future__ import annotations

from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.docx_formatter.domain.finding import Finding
from app.services.docx_formatter.domain.rule import AnalyzeRule
from app.services.docx_formatter.engine.context_classifier import (
    classify_paragraph_context,
    text_preview,
)
from app.services.docx_formatter.rules.common_format import CommonFormatMixin

TEXT_DECORATION_CONTEXTS = {
    "body_paragraph",
    "list_item",
    "heading",
    "chapter_number",
    "chapter_title",
    "caption",
}


class AdvancedFormatRule(CommonFormatMixin, AnalyzeRule):
    def analyze(self, doc: Any, config: dict[str, Any]) -> list[Finding]:
        advanced_config = config.get("advanced_review", {})
        if advanced_config.get("enabled", True) is False:
            return []

        findings: list[Finding] = []
        for paragraph_index, paragraph in enumerate(doc.paragraphs, start=1):
            context = classify_paragraph_context(paragraph, paragraph_index)
            findings.extend(
                self._paragraph_text_decoration_findings(
                    paragraph=paragraph,
                    paragraph_index=paragraph_index,
                    context=context,
                )
            )
            findings.extend(
                self._equation_findings(
                    paragraph=paragraph,
                    paragraph_index=paragraph_index,
                    context=context,
                )
            )

        findings.extend(self._table_text_decoration_findings(doc))
        return findings

    def _paragraph_text_decoration_findings(
        self,
        *,
        paragraph: Any,
        paragraph_index: int,
        context: Any,
    ) -> list[Finding]:
        if context.context not in TEXT_DECORATION_CONTEXTS:
            return []

        issues = _run_decoration_issues(paragraph.runs)
        if not issues:
            return []

        return [
            _manual_finding(
                finding_type="TEXT_DECORATION_REVIEW",
                severity="warning",
                location=f"Paragraph {paragraph_index}",
                message="Đoạn này có màu chữ, highlight hoặc gạch chân cần kiểm tra.",
                current_value="; ".join(issues),
                expected_value="Chữ nội dung nên dùng màu mặc định, không highlight và không gạch chân nếu mẫu không yêu cầu.",
                suggestion="Kiểm tra lại định dạng chữ tại đoạn này; chỉ giữ màu, highlight hoặc gạch chân khi đó là yêu cầu trình bày hợp lệ.",
                metadata={
                    "target": "paragraph",
                    "context": context.context,
                    "report_group_id": "text_decoration",
                    "report_severity": "major",
                    "paragraph_index": paragraph_index,
                    "field": "text_decoration",
                    "text_preview": context.text_preview,
                    "style_name": context.style_name,
                },
            )
        ]

    def _equation_findings(
        self,
        *,
        paragraph: Any,
        paragraph_index: int,
        context: Any,
    ) -> list[Finding]:
        if not _paragraph_contains_equation(paragraph):
            return []

        alignment = self._effective_alignment(paragraph)
        current_parts = ["Có công thức/ký hiệu toán học trong đoạn."]
        if alignment is None:
            current_parts.append("Không xác định được căn lề công thức.")
        elif alignment != WD_ALIGN_PARAGRAPH.CENTER:
            current_parts.append(f"Căn lề hiện tại: {self._alignment_label(alignment)}.")

        return [
            _manual_finding(
                finding_type="EQUATION_LAYOUT_REVIEW",
                severity="warning",
                location=f"Paragraph {paragraph_index}",
                message="Công thức/ký hiệu toán học cần kiểm tra bố cục thủ công.",
                current_value=" ".join(current_parts),
                expected_value=(
                    "Công thức trình bày riêng nên căn giữa, đánh số nhất quán và không bị cắt khi sang trang."
                ),
                suggestion=(
                    "Mở Word để kiểm tra căn giữa, số thứ tự công thức, font ký hiệu và vị trí ngắt trang quanh công thức."
                ),
                metadata={
                    "target": "paragraph",
                    "context": "equation",
                    "report_group_id": "equation_layout",
                    "report_severity": "major",
                    "paragraph_index": paragraph_index,
                    "field": "equation_layout",
                    "text_preview": context.text_preview,
                    "style_name": context.style_name,
                },
            )
        ]

    def _table_text_decoration_findings(self, doc: Any) -> list[Finding]:
        findings: list[Finding] = []
        for table_index, table in enumerate(doc.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                for cell_index, cell in enumerate(row.cells, start=1):
                    for cell_paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                        preview = text_preview(paragraph.text)
                        if not preview:
                            continue
                        issues = _run_decoration_issues(paragraph.runs)
                        if not issues:
                            continue
                        findings.append(
                            _manual_finding(
                                finding_type="TABLE_TEXT_DECORATION_REVIEW",
                                severity="warning",
                                location=(
                                    f"Table {table_index}, Row {row_index}, Cell {cell_index}, "
                                    f"Paragraph {cell_paragraph_index}"
                                ),
                                message="Chữ trong bảng có màu, highlight hoặc gạch chân cần kiểm tra.",
                                current_value="; ".join(issues),
                                expected_value=(
                                    "Chữ trong bảng nên dùng màu mặc định, không highlight và không gạch chân nếu mẫu không yêu cầu."
                                ),
                                suggestion=(
                                    "Kiểm tra lại định dạng chữ trong ô bảng; chỉ giữ màu, highlight hoặc gạch chân khi hợp lệ."
                                ),
                                metadata={
                                    "target": "table_cell",
                                    "context": "table_cell",
                                    "report_group_id": "text_decoration",
                                    "report_severity": "major",
                                    "table_index": table_index,
                                    "row_index": row_index,
                                    "cell_index": cell_index,
                                    "table_paragraph_index": cell_paragraph_index,
                                    "field": "text_decoration",
                                    "text_preview": preview,
                                    "style_name": getattr(getattr(paragraph, "style", None), "name", ""),
                                },
                            )
                        )
        return findings


def _run_decoration_issues(runs: list[Any]) -> list[str]:
    issues: list[str] = []
    seen: set[str] = set()

    for run in runs:
        if not run.text.strip():
            continue

        color = _run_color(run)
        if color and color.upper() != "000000":
            _append_once(issues, seen, f"Màu chữ khác mặc định: #{color.upper()}")

        highlight = getattr(run.font, "highlight_color", None)
        if highlight is not None:
            _append_once(issues, seen, f"Có highlight: {highlight}")

        underline = getattr(run.font, "underline", None)
        if underline:
            _append_once(issues, seen, "Có gạch chân")

    return issues


def _run_color(run: Any) -> str | None:
    try:
        rgb = run.font.color.rgb
    except AttributeError:
        return None
    if rgb is None:
        return None
    return str(rgb)


def _paragraph_contains_equation(paragraph: Any) -> bool:
    xml = getattr(getattr(paragraph, "_p", None), "xml", "")
    return "<m:oMath" in xml or "<m:oMathPara" in xml


def _append_once(items: list[str], seen: set[str], value: str) -> None:
    if value in seen:
        return
    seen.add(value)
    items.append(value)


def _manual_finding(
    *,
    finding_type: str,
    severity: str,
    location: str,
    message: str,
    current_value: str | None,
    expected_value: str | None,
    suggestion: str,
    metadata: dict[str, Any],
) -> Finding:
    return Finding(
        type=finding_type,
        severity=severity,
        location=location,
        message=message,
        current_value=current_value,
        expected_value=expected_value,
        suggestion=suggestion,
        metadata={
            **metadata,
            "auto_fixable": False,
            "manual_review": True,
            "fix_action": {
                "type": "manual_review",
                "reason": "Rule nâng cao này chỉ báo lỗi để người dùng kiểm tra thủ công, không tự sửa.",
            },
        },
    )
