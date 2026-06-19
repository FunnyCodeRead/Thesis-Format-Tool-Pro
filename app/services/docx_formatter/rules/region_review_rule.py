from __future__ import annotations

import re
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.docx_formatter.domain.finding import Finding
from app.services.docx_formatter.domain.rule import AnalyzeRule
from app.services.docx_formatter.engine.context_classifier import (
    classify_paragraph_context,
    text_preview,
)
from app.services.docx_formatter.rules.common_format import CommonFormatMixin
from app.services.docx_formatter.utils.docx_units import format_cm, length_to_cm
from app.services.docx_formatter.utils.text_utils import heading_key, normalize_text


class RegionReviewRule(CommonFormatMixin, AnalyzeRule):
    def analyze(self, doc: Any, config: dict[str, Any]) -> list[Finding]:
        findings: list[Finding] = []
        heading_path: list[str] = []
        paragraphs = list(doc.paragraphs)
        content_width_cm = _document_content_width_cm(doc)

        for paragraph_index, paragraph in enumerate(paragraphs, start=1):
            context = classify_paragraph_context(paragraph, paragraph_index)
            current_heading_key = heading_key(paragraph)
            if current_heading_key is not None and context.context == "heading":
                heading_path = _update_heading_path(
                    heading_path,
                    paragraph.text.strip(),
                    current_heading_key,
                )

            if context.context == "cover":
                findings.extend(self._cover_findings(paragraph, paragraph_index, context))
            elif context.context in {"toc", "list_of_figures", "list_of_tables"}:
                findings.extend(
                    self._list_findings(
                        paragraph=paragraph,
                        paragraph_index=paragraph_index,
                        context=context,
                        heading_path=heading_path,
                    )
                )
            elif context.context == "caption":
                findings.extend(
                    self._caption_findings(
                        paragraph=paragraph,
                        paragraph_index=paragraph_index,
                        context=context,
                        config=config,
                        heading_path=heading_path,
                    )
                )

            if (
                self._paragraph_contains_image(paragraph)
                and not self._should_skip_image_review(context.context)
                and not self._has_nearby_caption(paragraphs, paragraph_index)
            ):
                findings.append(
                    _finding(
                        finding_type="IMAGE_LAYOUT_REVIEW",
                        severity="warning",
                        location=f"Paragraph {paragraph_index}",
                        message="Ảnh hoặc sơ đồ trong nội dung có thể thiếu caption gần đó.",
                        current_value=context.text_preview or "Ảnh/sơ đồ trong đoạn văn",
                        expected_value="Ảnh/sơ đồ trong phần nội dung chính cần có caption gần vị trí ảnh.",
                        suggestion="Kiểm tra ảnh/sơ đồ trong nội dung và bổ sung hoặc cập nhật caption gần ảnh nếu cần.",
                        metadata={
                            "target": "paragraph",
                            "context": "image_layout",
                            "report_group_id": "image_layout",
                            "report_severity": "major",
                            "paragraph_index": paragraph_index,
                            "heading_path": heading_path.copy(),
                            "field": "image_layout",
                            "text_preview": context.text_preview,
                            "style_name": context.style_name,
                        },
                    )
                )

            if self._paragraph_contains_image(paragraph) and not self._should_skip_image_review(context.context):
                findings.extend(
                    self._image_layout_findings(
                        paragraph=paragraph,
                        paragraph_index=paragraph_index,
                        context=context,
                        heading_path=heading_path,
                        content_width_cm=content_width_cm,
                    )
                )

        findings.extend(self._table_findings(doc, config))
        findings.extend(self._header_footer_findings(doc, config))
        return findings

    def _cover_findings(
        self,
        paragraph: Any,
        paragraph_index: int,
        context: Any,
    ) -> list[Finding]:
        findings: list[Finding] = []
        alignment = self._effective_alignment(paragraph)
        if alignment is not None and alignment != WD_ALIGN_PARAGRAPH.CENTER:
            findings.append(
                _finding(
                    finding_type="COVER_ALIGNMENT_REVIEW",
                    severity="warning",
                    location=f"Paragraph {paragraph_index}",
                    message="Một dòng trên trang bìa chưa được căn giữa.",
                    current_value=self._alignment_label(alignment),
                    expected_value="CENTER",
                    suggestion="Kiểm tra bố cục trang bìa và căn giữa các khối chữ theo mẫu trường.",
                    metadata={
                        "target": "paragraph",
                        "context": "cover",
                        "report_group_id": "cover_layout",
                        "report_severity": "critical",
                        "paragraph_index": paragraph_index,
                        "field": "alignment",
                        "text_preview": context.text_preview,
                        "style_name": context.style_name,
                    },
                )
            )
        return findings

    def _list_findings(
        self,
        *,
        paragraph: Any,
        paragraph_index: int,
        context: Any,
        heading_path: list[str],
    ) -> list[Finding]:
        text = paragraph.text.strip()
        normalized_text = normalize_text(text)
        findings: list[Finding] = []
        group_id = context.context

        alignment = self._effective_alignment(paragraph)
        if _is_list_title(normalized_text) and alignment is not None and alignment != WD_ALIGN_PARAGRAPH.CENTER:
            findings.append(
                _finding(
                    finding_type=f"{group_id.upper()}_TITLE_ALIGNMENT_REVIEW",
                    severity="warning",
                    location=f"Paragraph {paragraph_index}",
                    message="Tiêu đề danh mục/mục lục chưa được căn giữa.",
                    current_value=self._alignment_label(alignment),
                    expected_value="CENTER",
                    suggestion="Căn giữa tiêu đề danh mục/mục lục theo mẫu trường.",
                    metadata={
                        "target": "paragraph",
                        "context": context.context,
                        "report_group_id": group_id,
                        "report_severity": "major",
                        "paragraph_index": paragraph_index,
                        "heading_path": heading_path.copy(),
                        "field": "alignment",
                        "text_preview": context.text_preview,
                        "style_name": context.style_name,
                    },
                )
            )

        if _looks_like_list_entry(normalized_text) and not re.search(r"\.{3,}", text):
            findings.append(
                _finding(
                    finding_type=f"{group_id.upper()}_DOT_LEADER_REVIEW",
                    severity="warning",
                    location=f"Paragraph {paragraph_index}",
                    message="Dòng danh mục có thể thiếu dot leader hoặc số trang căn phải.",
                    current_value=context.text_preview,
                    expected_value="Dòng danh mục nên dùng dot leader và số trang căn phải.",
                    suggestion="Cập nhật hoặc tạo lại danh mục trong Word rồi kiểm tra căn phải số trang.",
                    metadata={
                        "target": "paragraph",
                        "context": context.context,
                        "report_group_id": group_id,
                        "report_severity": "major",
                        "paragraph_index": paragraph_index,
                        "heading_path": heading_path.copy(),
                        "field": "list_entry_layout",
                        "text_preview": context.text_preview,
                        "style_name": context.style_name,
                    },
                )
            )
        return findings

    def _caption_findings(
        self,
        *,
        paragraph: Any,
        paragraph_index: int,
        context: Any,
        config: dict[str, Any],
        heading_path: list[str],
    ) -> list[Finding]:
        expected = {
            "font_name": config.get("paragraph", {}).get("font_name", "Times New Roman"),
            "font_size": config.get("paragraph", {}).get("font_size", 13),
            "alignment": "CENTER",
            "space_before_pt": 6,
            "space_after_pt": 6,
        }
        findings = self.analyze_common_format(
            paragraph=paragraph,
            expected=expected,
            location=f"Paragraph {paragraph_index}",
            type_prefix="CAPTION",
            metadata={
                "target": "caption",
                "context": "caption",
                "report_group_id": "caption",
                "report_severity": "minor",
                "auto_fixable": True,
                "manual_review": False,
                "paragraph_index": paragraph_index,
                "heading_path": heading_path.copy(),
                "text_preview": context.text_preview,
                "style_name": context.style_name,
            },
        )
        return findings

    def _table_findings(self, doc: Any, config: dict[str, Any]) -> list[Finding]:
        expected = {
            "font_name": config.get("paragraph", {}).get("font_name", "Times New Roman"),
            "font_size": config.get("paragraph", {}).get("font_size", 13),
        }
        findings: list[Finding] = []
        content_width_cm = _document_content_width_cm(doc)
        for table_index, table in enumerate(doc.tables, start=1):
            table_width_cm = _table_width_cm(table)
            if (
                content_width_cm is not None
                and table_width_cm is not None
                and table_width_cm > content_width_cm + 0.2
            ):
                findings.append(
                    _finding(
                        finding_type="TABLE_WIDTH_REVIEW",
                        severity="warning",
                        location=f"Table {table_index}",
                        message="Bảng có thể rộng hơn vùng nội dung của trang.",
                        current_value=format_cm(table_width_cm),
                        expected_value=f"Không vượt quá vùng nội dung {format_cm(content_width_cm)}",
                        suggestion="Kiểm tra lại độ rộng bảng, có thể cần chỉnh column width hoặc đưa bảng lớn sang phụ lục/section ngang.",
                        metadata={
                            "target": "table_cell",
                            "context": "table_cell",
                            "report_group_id": "table",
                            "report_severity": "major",
                            "table_index": table_index,
                            "row_index": 1,
                            "cell_index": 1,
                            "table_paragraph_index": 1,
                            "field": "table_width",
                            "text_preview": "Table width review",
                            "style_name": "Table",
                        },
                    )
                )

        return findings

    def _header_footer_findings(self, doc: Any, config: dict[str, Any]) -> list[Finding]:
        findings: list[Finding] = []
        expected = {
            "font_name": config.get("paragraph", {}).get("font_name", "Times New Roman"),
            "font_size": config.get("paragraph", {}).get("font_size", 13),
        }
        for section_index, section in enumerate(doc.sections, start=1):
            footer_text = _part_text(section.footer)
            if not _contains_page_number_field(section.footer):
                findings.append(
                    _finding(
                        finding_type="PAGE_NUMBER_FOOTER_REVIEW",
                        severity="warning",
                        location=f"Section {section_index} Footer",
                        message="Footer/số trang cần được kiểm tra.",
                        current_value=footer_text or "Footer trống hoặc chưa có PAGE field hiển thị.",
                        expected_value="Số trang cần được cấu hình trong footer theo từng section.",
                        suggestion="Kiểm tra footer, số trang La Mã ở phần đầu và số trang Ả Rập ở phần nội dung chính.",
                        metadata={
                            "target": "footer",
                            "context": "header_footer",
                            "report_group_id": "header_footer_page_number",
                            "report_severity": "major",
                            "section_index": section_index,
                            "field": "page_number_position",
                            "text_preview": footer_text,
                            "style_name": "Footer",
                        },
                    )
                )

            header_text = _part_text(section.header)
            if header_text and _contains_body_like_page_number(header_text):
                findings.append(
                    _finding(
                        finding_type="HEADER_FOOTER_LAYOUT_REVIEW",
                        severity="warning",
                        location=f"Section {section_index} Header",
                        message="Header/footer có thể chứa số trang hoặc chữ bị đặt sai vị trí.",
                        current_value=header_text,
                        expected_value="Header/footer chỉ nên chứa nội dung chạy đầu trang/chân trang và field số trang đúng vị trí.",
                        suggestion="Kiểm tra field header/footer; nếu có chữ từ thân bài bị chuyển nhầm thì xử lý thủ công trong Word.",
                        metadata={
                            "target": "header",
                            "context": "header_footer",
                            "report_group_id": "header_footer_page_number",
                            "report_severity": "major",
                            "section_index": section_index,
                            "field": "header_footer_layout",
                            "text_preview": header_text,
                            "style_name": "Header",
                        },
                    )
                )
        return findings

    def _image_layout_findings(
        self,
        *,
        paragraph: Any,
        paragraph_index: int,
        context: Any,
        heading_path: list[str],
        content_width_cm: float | None,
    ) -> list[Finding]:
        findings: list[Finding] = []
        if self._is_floating_image(paragraph):
            findings.append(
                _finding(
                    finding_type="IMAGE_WRAP_REVIEW",
                    severity="warning",
                    location=f"Paragraph {paragraph_index}",
                    message="Ảnh/sơ đồ đang dùng floating hoặc wrap text cần kiểm tra.",
                    current_value="Ảnh đang dùng kiểu floating/wrap text",
                    expected_value="Ảnh trong nội dung nên có wrapping nhất quán và không che lên chữ.",
                    suggestion="Kiểm tra tùy chọn bố cục của ảnh; ưu tiên đặt ảnh nằm cùng dòng với chữ nếu mẫu trường yêu cầu.",
                    metadata={
                        "target": "paragraph",
                        "context": "image_layout",
                        "report_group_id": "image_layout",
                        "report_severity": "major",
                        "paragraph_index": paragraph_index,
                        "heading_path": heading_path.copy(),
                        "field": "image_wrap",
                        "text_preview": context.text_preview,
                        "style_name": context.style_name,
                    },
                )
            )

        if content_width_cm is None:
            return findings

        for image_width_cm in self._image_widths_cm(paragraph):
            if image_width_cm <= content_width_cm + 0.2:
                continue
            findings.append(
                _finding(
                    finding_type="IMAGE_WIDTH_REVIEW",
                    severity="warning",
                    location=f"Paragraph {paragraph_index}",
                    message="Ảnh/sơ đồ có thể rộng hơn vùng nội dung.",
                    current_value=format_cm(image_width_cm),
                    expected_value=f"Không vượt quá vùng nội dung {format_cm(content_width_cm)}",
                    suggestion="Kiểm tra kích thước ảnh/sơ đồ và thu nhỏ nếu ảnh tràn lề.",
                    metadata={
                        "target": "paragraph",
                        "context": "image_layout",
                        "report_group_id": "image_layout",
                        "report_severity": "major",
                        "paragraph_index": paragraph_index,
                        "heading_path": heading_path.copy(),
                        "field": "image_width",
                        "text_preview": context.text_preview,
                        "style_name": context.style_name,
                    },
                )
            )
            break

        return findings

    def _header_footer_format_findings(
        self,
        *,
        part: Any,
        expected: dict[str, Any],
        section_index: int,
        target: str,
        style_name: str,
    ) -> list[Finding]:
        findings: list[Finding] = []
        for part_paragraph_index, paragraph in enumerate(part.paragraphs, start=1):
            preview = text_preview(paragraph.text)
            if not preview:
                continue
            part_findings = self.analyze_common_format(
                paragraph=paragraph,
                expected=expected,
                location=f"Section {section_index} {style_name}",
                type_prefix="HEADER_FOOTER",
                metadata={
                    "target": target,
                    "context": "header_footer",
                    "report_group_id": "header_footer_page_number",
                    "report_severity": "minor",
                    "auto_fixable": False,
                    "manual_review": True,
                    "section_index": section_index,
                    "part_paragraph_index": part_paragraph_index,
                    "text_preview": preview,
                    "style_name": style_name,
                },
            )
            findings.extend(_with_manual_review(finding) for finding in part_findings)
        return findings

    def _is_floating_image(self, paragraph: Any) -> bool:
        xml = getattr(getattr(paragraph, "_p", None), "xml", "")
        return "<wp:anchor" in xml

    def _image_widths_cm(self, paragraph: Any) -> list[float]:
        xml = getattr(getattr(paragraph, "_p", None), "xml", "")
        widths: list[float] = []
        for value in re.findall(r"<wp:extent[^>]*cx=\"(\d+)\"", xml):
            width_cm = _emu_to_cm(value)
            if width_cm is not None:
                widths.append(width_cm)
        return widths

    def _paragraph_contains_image(self, paragraph: Any) -> bool:
        xml = getattr(getattr(paragraph, "_p", None), "xml", "")
        return "<w:drawing" in xml or "<w:pict" in xml

    def _should_skip_image_review(self, context: str) -> bool:
        return context in {
            "cover",
            "front_matter_heading",
            "toc",
            "list_of_figures",
            "list_of_tables",
            "caption",
            "header_footer",
        }

    def _has_nearby_caption(self, paragraphs: list[Any], paragraph_index: int) -> bool:
        start = max(1, paragraph_index - 2)
        end = min(len(paragraphs), paragraph_index + 2)
        for nearby_index in range(start, end + 1):
            if nearby_index == paragraph_index:
                continue
            nearby = paragraphs[nearby_index - 1]
            nearby_context = classify_paragraph_context(nearby, nearby_index)
            if nearby_context.context == "caption":
                return True
        return False


def _finding(
    *,
    finding_type: str,
    severity: str,
    location: str,
    message: str,
    current_value: str | None,
    expected_value: str | None,
    suggestion: str,
    metadata: dict[str, Any],
) -> Finding:
    metadata = {
        **metadata,
        "auto_fixable": False,
        "manual_review": True,
        "fix_action": {
            "type": "manual_review",
            "reason": "Rule này chỉ phân tích và cần người dùng kiểm tra thủ công.",
        },
    }
    return Finding(
        type=finding_type,
        severity=severity,
        location=location,
        message=message,
        current_value=current_value,
        expected_value=expected_value,
        suggestion=suggestion,
        metadata=metadata,
    )


def _document_content_width_cm(doc: Any) -> float | None:
    try:
        section = doc.sections[0]
    except (IndexError, TypeError):
        return None
    page_width_cm = length_to_cm(section.page_width)
    left_margin_cm = length_to_cm(section.left_margin) or 0
    right_margin_cm = length_to_cm(section.right_margin) or 0
    if page_width_cm is None:
        return None
    return page_width_cm - left_margin_cm - right_margin_cm


def _table_width_cm(table: Any) -> float | None:
    xml = getattr(getattr(table, "_tbl", None), "xml", "")
    grid_widths = [
        _twips_to_cm(value)
        for value in re.findall(r"<w:gridCol[^>]*w:w=\"(\d+)\"", xml)
    ]
    known_widths = [value for value in grid_widths if value is not None]
    if known_widths:
        return sum(known_widths)

    match = re.search(r"<w:tblW[^>]*w:type=\"dxa\"[^>]*w:w=\"(\d+)\"", xml)
    if match:
        return _twips_to_cm(match.group(1))
    match = re.search(r"<w:tblW[^>]*w:w=\"(\d+)\"[^>]*w:type=\"dxa\"", xml)
    if match:
        return _twips_to_cm(match.group(1))
    return None


def _twips_to_cm(value: str) -> float | None:
    try:
        return int(value) / 1440 * 2.54
    except (TypeError, ValueError):
        return None


def _emu_to_cm(value: str) -> float | None:
    try:
        return int(value) / 360000
    except (TypeError, ValueError):
        return None


def _with_manual_review(finding: Finding) -> Finding:
    metadata = {
        **finding.metadata,
        "auto_fixable": False,
        "manual_review": True,
        "fix_action": {
            "type": "manual_review",
            "reason": "Rule này chỉ phân tích và cần người dùng kiểm tra thủ công.",
        },
    }
    return Finding(
        type=finding.type,
        severity=finding.severity,
        location=finding.location,
        message=finding.message,
        current_value=finding.current_value,
        expected_value=finding.expected_value,
        suggestion=finding.suggestion,
        metadata=metadata,
    )


def _update_heading_path(current: list[str], text: str, key: str) -> list[str]:
    if key == "heading_1":
        return [text]
    if key == "heading_2":
        return [*(current[:1] or []), text]
    if key == "heading_3":
        return [*(current[:2] or []), text]
    return current


def _looks_like_list_entry(normalized_text: str) -> bool:
    return bool(re.search(r"\s\d+\s*$", normalized_text)) and len(normalized_text) > 12


def _is_list_title(normalized_text: str) -> bool:
    return normalized_text.startswith(
        (
            "MUC LUC",
            "TABLE OF CONTENTS",
            "DANH MUC HINH",
            "DANH SACH HINH",
            "LIST OF FIGURES",
            "DANH MUC BANG",
            "DANH SACH BANG",
            "LIST OF TABLES",
        )
    )


def _part_text(part: Any) -> str:
    text = " ".join(paragraph.text.strip() for paragraph in part.paragraphs if paragraph.text.strip())
    return text_preview(text)


def _contains_page_number_field(part: Any) -> bool:
    xml = " ".join(getattr(getattr(paragraph, "_p", None), "xml", "") for paragraph in part.paragraphs)
    return "PAGE" in xml.upper() or "NUMPAGES" in xml.upper()


def _contains_body_like_page_number(text: str) -> bool:
    normalized = normalize_text(text)
    return bool(re.search(r"\b\d+(?:IV|V|VI|VII|VIII|IX|X)\b", normalized))
