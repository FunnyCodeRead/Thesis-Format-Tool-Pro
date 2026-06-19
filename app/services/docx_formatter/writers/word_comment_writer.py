from __future__ import annotations

from dataclasses import asdict
from pathlib import Path
from typing import Any

from docx import Document

from app.services.docx_formatter.domain.annotation import AnnotationComment, AnnotationTarget

COMMENT_AUTHOR = "Thesis Format Tool Pro"
COMMENT_INITIALS = "TFTP"
NEARBY_PARAGRAPH_ANCHOR_WINDOW = 15


class WordCommentWriterError(RuntimeError):
    pass


class WordCommentWriter:
    def write(
        self,
        *,
        input_path: str,
        output_path: str,
        comments: list[AnnotationComment],
    ) -> dict[str, Any]:
        try:
            doc = Document(input_path)
        except Exception as exc:
            raise WordCommentWriterError("Failed to load .docx document.") from exc

        if not hasattr(doc, "add_comment"):
            return self._write_summary_fallback(doc, output_path, comments)

        total_findings = sum(comment.source_count for comment in comments)
        grouped_findings_in_comments = _grouped_findings_in_comments(comments)
        comment_count = 0
        skipped_reasons: list[dict[str, Any]] = []

        for comment in comments:
            runs = self._anchor_runs(doc, comment)
            if not runs:
                skipped_reasons.append(
                    self._skip_reason(
                        comment,
                        "Không tìm thấy đoạn/run Word phù hợp với locator của finding.",
                    )
                )
                continue

            try:
                doc.add_comment(
                    runs,
                    text=comment.to_text(),
                    author=COMMENT_AUTHOR,
                    initials=COMMENT_INITIALS,
                )
            except Exception as exc:
                raise WordCommentWriterError("Failed to add Word comment.") from exc
            comment_count += 1

        self._save(doc, output_path)
        skipped_findings = sum(item["finding_count"] for item in skipped_reasons)
        comment_note = self._comment_note(
            total_findings=total_findings,
            total_comments_created=comment_count,
            skipped_findings=skipped_findings,
        )
        return {
            "total_findings": total_findings,
            "total_comments_created": comment_count,
            "comment_count": comment_count,
            "comment_strategy": "hybrid",
            "grouped_findings_in_comments": grouped_findings_in_comments,
            "skipped_comments": len(skipped_reasons),
            "skipped_count": len(skipped_reasons),
            "skipped_findings": skipped_findings,
            "skipped_reason": skipped_reasons,
            "comment_note": comment_note,
            "fallback_summary_pages": 0,
            "file_size_bytes": Path(output_path).stat().st_size,
        }

    def _write_summary_fallback(
        self,
        doc: Any,
        output_path: str,
        comments: list[AnnotationComment],
    ) -> dict[str, Any]:
        doc.add_page_break()
        doc.add_heading("Báo cáo lỗi định dạng", level=1)
        doc.add_paragraph(
            "Runtime hiện tại chưa hỗ trợ Word comment trực tiếp, nên bản báo cáo này "
            "thêm trang tổng hợp lỗi ở cuối file."
        )

        for index, comment in enumerate(comments, start=1):
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"{index}. {comment.title}").bold = True
            doc.add_paragraph(comment.to_text())

        self._save(doc, output_path)
        total_findings = sum(comment.source_count for comment in comments)
        grouped_findings_in_comments = _grouped_findings_in_comments(comments)
        skipped_reasons = [
            self._skip_reason(
                comment,
                (
                    "Runtime python-docx không hỗ trợ Document.add_comment; "
                    "đã ghi lỗi vào trang tổng hợp thay vì Word comment thật."
                ),
            )
            for comment in comments
        ]
        return {
            "total_findings": total_findings,
            "total_comments_created": 0,
            "comment_count": 0,
            "comment_strategy": "hybrid",
            "grouped_findings_in_comments": grouped_findings_in_comments,
            "skipped_comments": len(skipped_reasons),
            "skipped_count": len(skipped_reasons),
            "skipped_findings": total_findings,
            "skipped_reason": skipped_reasons,
            "comment_note": self._comment_note(
                total_findings=total_findings,
                total_comments_created=0,
                skipped_findings=total_findings,
            ),
            "fallback_summary_pages": 1,
            "file_size_bytes": Path(output_path).stat().st_size,
        }

    def _anchor_runs(self, doc: Any, comment: AnnotationComment) -> list[Any]:
        target = comment.target

        if target.target in {"paragraph", "heading", "caption"}:
            return self._runs_for_paragraph_index(doc, target.paragraph_index)

        if target.target == "table_cell":
            runs = self._runs_for_table_cell(doc, target)
            if runs:
                return runs
            return self._runs_for_paragraph_index(doc, target.paragraph_index)

        if target.target == "header":
            runs = self._runs_for_header_or_footer(doc, target, part_name="header")
            if runs:
                return runs
            return self._runs_for_section_start(doc, target.section_index)

        if target.target == "footer":
            runs = self._runs_for_header_or_footer(doc, target, part_name="footer")
            if runs:
                return runs
            if target.paragraph_index is not None:
                runs = self._runs_for_paragraph_index(doc, target.paragraph_index)
                if runs:
                    return runs
            return self._runs_for_section_start(doc, target.section_index)

        if target.target == "section":
            return self._runs_for_section_start(doc, target.section_index)

        return []

    def _runs_for_paragraph_index(self, doc: Any, paragraph_index: int | None) -> list[Any]:
        if paragraph_index is None:
            return []

        zero_based_index = paragraph_index - 1
        if zero_based_index < 0 or zero_based_index >= len(doc.paragraphs):
            return []

        runs = self._commentable_runs(doc.paragraphs[zero_based_index])
        if runs:
            return runs

        return self._runs_for_nearby_paragraph_anchor(doc, zero_based_index)

    def _runs_for_nearby_paragraph_anchor(self, doc: Any, zero_based_index: int) -> list[Any]:
        for offset in range(1, NEARBY_PARAGRAPH_ANCHOR_WINDOW + 1):
            forward_index = zero_based_index + offset
            if forward_index < len(doc.paragraphs):
                runs = self._commentable_runs(doc.paragraphs[forward_index])
                if runs:
                    return runs

            backward_index = zero_based_index - offset
            if backward_index >= 0:
                runs = self._commentable_runs(doc.paragraphs[backward_index])
                if runs:
                    return runs

        return []

    def _runs_for_table_cell(self, doc: Any, target: AnnotationTarget) -> list[Any]:
        if target.table_index is None or target.row_index is None or target.cell_index is None:
            return []

        table_index = target.table_index - 1
        row_index = target.row_index - 1
        cell_index = target.cell_index - 1

        if table_index < 0 or table_index >= len(doc.tables):
            return []

        table = doc.tables[table_index]
        if row_index < 0 or row_index >= len(table.rows):
            return []

        cells = table.rows[row_index].cells
        if cell_index < 0 or cell_index >= len(cells):
            return self._runs_for_first_table_anchor(table)

        target_paragraph_index = target.table_paragraph_index
        if target_paragraph_index is not None:
            paragraph_index = target_paragraph_index - 1
            if paragraph_index < 0 or paragraph_index >= len(cells[cell_index].paragraphs):
                return self._runs_for_first_table_anchor(table)
            runs = self._commentable_runs(cells[cell_index].paragraphs[paragraph_index])
            if runs:
                return runs
            return self._runs_for_first_table_anchor(table)

        for paragraph in cells[cell_index].paragraphs:
            runs = self._commentable_runs(paragraph)
            if runs:
                return runs
        return self._runs_for_first_table_anchor(table)

    def _runs_for_first_table_anchor(self, table: Any) -> list[Any]:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    runs = self._commentable_runs(paragraph)
                    if runs:
                        return runs
        return []

    def _runs_for_header_or_footer(
        self,
        doc: Any,
        target: AnnotationTarget,
        *,
        part_name: str,
    ) -> list[Any]:
        section = self._section(doc, target.section_index)
        if section is None:
            return []

        part = getattr(section, part_name, None)
        if part is None:
            return []

        if target.part_paragraph_index is not None:
            paragraph_index = target.part_paragraph_index - 1
            if 0 <= paragraph_index < len(part.paragraphs):
                runs = self._commentable_runs(part.paragraphs[paragraph_index])
                if runs:
                    return runs

        for paragraph in part.paragraphs:
            runs = self._commentable_runs(paragraph)
            if runs:
                return runs
        return []

    def _runs_for_section_start(self, doc: Any, section_index: int | None) -> list[Any]:
        paragraph_index = self._section_start_paragraph_index(doc, section_index)
        if paragraph_index is None:
            return []

        for candidate_index in range(paragraph_index, len(doc.paragraphs) + 1):
            runs = self._runs_for_paragraph_index(doc, candidate_index)
            if runs:
                return runs
        return []

    def _section_start_paragraph_index(self, doc: Any, section_index: int | None) -> int | None:
        if section_index is None:
            return None
        if section_index <= 1:
            return 1
        if section_index > len(doc.sections):
            return None

        section_breaks_seen = 1
        for index, paragraph in enumerate(doc.paragraphs, start=1):
            if self._paragraph_has_section_properties(paragraph):
                section_breaks_seen += 1
                if section_breaks_seen == section_index:
                    return min(index + 1, len(doc.paragraphs))

        return None

    def _section(self, doc: Any, section_index: int | None) -> Any | None:
        if section_index is None:
            section_index = 1
        zero_based_index = section_index - 1
        if zero_based_index < 0 or zero_based_index >= len(doc.sections):
            return None
        return doc.sections[zero_based_index]

    def _paragraph_has_section_properties(self, paragraph: Any) -> bool:
        p_pr = getattr(paragraph._p, "pPr", None)
        return p_pr is not None and getattr(p_pr, "sectPr", None) is not None

    def _text_runs(self, paragraph: Any) -> list[Any]:
        runs = [run for run in paragraph.runs if run.text and run.text.strip()]
        if runs:
            return runs

        if not hasattr(paragraph, "iter_inner_content"):
            return []

        inner_runs: list[Any] = []
        for item in paragraph.iter_inner_content():
            item_runs = getattr(item, "runs", None)
            if not item_runs:
                continue
            inner_runs.extend(run for run in item_runs if run.text and run.text.strip())
        return inner_runs

    def _commentable_runs(self, paragraph: Any) -> list[Any]:
        text_runs = self._text_runs(paragraph)
        if text_runs:
            return text_runs

        return [
            run
            for run in paragraph.runs
            if "<w:drawing" in getattr(run._r, "xml", "") or "<w:pict" in getattr(run._r, "xml", "")
        ]

    def _skip_reason(self, comment: AnnotationComment, reason: str) -> dict[str, Any]:
        return {
            "reason": reason,
            "finding_count": comment.source_count,
            "finding_ids": comment.source_ids,
            "source_types": [issue.source_type for issue in comment.issues],
            "target": asdict(comment.target),
            "title": comment.title,
        }

    def _comment_note(
        self,
        *,
        total_findings: int,
        total_comments_created: int,
        skipped_findings: int,
    ) -> str:
        if skipped_findings:
            return (
                f"Có {skipped_findings} finding không tạo được comment vì không tìm thấy "
                "anchor phù hợp; xem skipped_reason để biết từng lỗi bị bỏ qua."
            )
        if total_findings > total_comments_created:
            return (
                "Không bỏ qua finding nào; chênh lệch là do chế độ hybrid gom nhiều lỗi "
                "cùng style hoặc cùng paragraph/element vào một comment đại diện."
            )
        return "Tất cả finding đã được tạo comment."

    def _save(self, doc: Any, output_path: str) -> None:
        try:
            doc.save(output_path)
        except Exception as exc:
            raise WordCommentWriterError("Failed to save annotated .docx document.") from exc


def _grouped_findings_in_comments(comments: list[AnnotationComment]) -> int:
    return sum(max(0, comment.source_count - 1) for comment in comments)
